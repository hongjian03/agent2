import streamlit as st
import json
from langchain_openai import ChatOpenAI
from langchain.prompts import ChatPromptTemplate
from langchain.schema import HumanMessage, SystemMessage
from langchain.callbacks.streaming_stdout import StreamingStdOutCallbackHandler
from langchain.callbacks.streamlit import StreamlitCallbackHandler
from langchain.chains import SequentialChain, LLMChain
import os
from typing import Dict, Any, List
import logging
import sys
from docx import Document
import io
import base64
from PyPDF2 import PdfReader
from PIL import Image
import fitz  # PyMuPDF
# 配置日志记录
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(levelname)s - %(message)s',
    handlers=[
        logging.StreamHandler(sys.stdout)
    ]
)
from queue import Queue
from threading import Thread
import time
from queue import Empty
logger = logging.getLogger(__name__)
from langchain.callbacks.base import BaseCallbackHandler
import markitdown

# 记录程序启动
logger.info("程序开始运行")

# 只在第一次运行时替换 sqlite3
if 'sqlite_setup_done' not in st.session_state:
    try:
        logger.info("尝试设置 SQLite")
        __import__('pysqlite3')
        sys.modules['sqlite3'] = sys.modules.pop('pysqlite3')
        st.session_state.sqlite_setup_done = True
        logger.info("SQLite 设置成功")
    except Exception as e:
        logger.error(f"SQLite 设置错误: {str(e)}")
        st.session_state.sqlite_setup_done = True


class PromptTemplates:
    def __init__(self):
        # 定义示例数据作为字符串
        self.default_templates = {
            'transcript_role': """
            # 角色
            你是专业的成绩单分析师，擅长从成绩单中提取关键信息并以表格形式展示成绩。
            """,
            
            'transcript_task': """

            """,
            
            'transcript_output': """

            """,
            
            'consultant_role2': """
            # 角色
            我是一位资深的学术内容创作专家，曾为数百名学生成功撰写录取率极高的个人陈述。
            我拥有丰富的跨文化写作经验，精通中英双语表达，能够在保持文化真实性的同时创作符合英语思维和表达习惯的内容。
            我曾在多所国际顶尖大学的招生部门工作，深谙各学科领域的录取偏好和评判标准。
            我擅长将学生的素材转化为引人入胜的叙事，通过巧妙的结构安排和语言选择，使每位申请者的独特价值在众多申请中脱颖而出。

            """,
            
            'output_format2': """
            输出格式：
            输出内容为字符串，不是代码块，开头严禁添加"点点点markdown"这样的代码块标识，注意是严禁添加
            # 个人陈述（专业大类1：[专业名称]）

            ## 主要段落结构

            ### 1. 专业兴趣塑造
            [按照分析报告中的段落素材策略与增强指南组织内容，注重逻辑性，并且深入展开细节描述和观点叙述，减少素材的堆砌，注重描述的深度...]

            ### 2. 学术基础展示
            [按照分析报告中的段落素材策略与增强指南组织内容，突出相关课程的学习成果和技能提升，体现与该专业方向的契合...]

            ### 3. 研究经历深化
            [按照分析报告中的段落素材策略与增强指南组织内容，遵循STAR原则和总分总结构详细描述相关经历，与专业方向相联系...]

            ### 4. 实习经历深化
            [按照分析报告中的段落素材策略与增强指南组织内容，遵循STAR原则和总分总结构详细描述相关经历，与专业方向相联系...]

            ### 5. 未来规划提升
            [按照分析报告中的段落素材策略与增强指南组织内容，结合该专业方向提供具体且合理的规划...]

            ### 6. 为何选择该专业和院校
            [从专业发展前景、学术理念契合和个人成长等角度，逻辑性地阐述为何选择该专业领域深造，不针对具体学校]

            ### 7. 结语
            [简洁有力地总结申请者的优势、志向和对该专业的热情...]

            ---

            # 个人陈述（专业大类2：[专业名称]）
            [按相同结构组织第二个专业方向的个人陈述...]

            ---

            # 写作说明

            ## 格式与结构要求
            - 确保文章结构清晰，段落之间有良好的逻辑过渡
            - 所有非素材表中需要补充的内容必须保留中文并用【补充：】标记
            - 内容均使用纯中文表达，表述则符合英语思维习惯
            - 确保英文部分与中文补充内容之间的自然过渡
            - 技术术语和专业概念则使用准确的英文表达

            ## 内容质量要求
            - 保持文章的整体连贯性和专业性
            - 重点突出申请者的优势，并与申请方向建立明确联系
            - 内容应真实可信，避免虚构经历或夸大成就
            - 每个主题部分应当是一个连贯的整体段落，而非多个松散段落

            ## 专业相关要求
            - "为何选择该专业和院校"部分应从专业角度进行逻辑论述，不针对具体学校
            - 在分析和撰写过程中充分参考成绩单中的具体课程表现，但不要体现任何分数

            ## 写作风格要求
            - 确保内容精练，避免不必要的重复和冗余表达
            - 结语应简明扼要地总结全文，展现申请者的决心和愿景
            - 所有段落必须严格执行报告中的指导，不合并不同经历或添加额外内容

            ## 逻辑连贯性要求
            - 避免出现"突然感兴趣"或"因此更感兴趣"等生硬转折，确保兴趣发展有合理的渐进过程
            - 各段落间应有内在的逻辑联系，而非简单罗列，每段内容应自然引出下一段内容
            - 确保经历与专业兴趣间的关联性具有说服力，展示清晰的思维发展路径

            ## 整体要求
            - 在严格遵循分析报告的基础上，优化表述逻辑，确保内容流畅自然，避免生硬的连接和突兀的转折
            - 即使是按照报告中的指导进行写作，也应确保整体叙事具有内在一致性和合理的心理动机发展


            """,
            
            'consultant_task2': """
            任务描述
            1.基于提供的素材表、申请方向、成绩单、定制需求，为指定的专业方向创作完整的个人陈述初稿
            2.严格按照分析报告中的"段落素材策略与增强指南"组织内容，确保每个段落都完全遵循报告中提供的所有个性化需求及增强策略点
            3.遵循STAR原则(情境-任务-行动-结果)呈现研究经历和实习经历，且只选择报告中指定的一个最相关经历
            4.突出申请者与申请方向的契合点，参考"招生倾向分析"部分的具体要点
            5.在正文中直接使用【补充：】标记所有非素材表中的内容
            6.确保段落间有自然过渡，保持文章整体连贯性
            7.所有段落内容必须严格遵循报告中对应的策略要点和增强策略，不添加未在报告中指定的内容
            8.科研经历和实习经历段落必须各自只选择一个最相关的经历，完全按照报告中指定的增强策略展开
            9.在遵循分析报告内容的同时，需要优化表述逻辑，确保内容之间的连贯性和自然过渡

            结构要求
            1.每个段落应采用"总-分-总"结构，第一句话承上启下，最后一句话总结该经历与目标专业的联系
            2.每个段落应有连贯性，作为一个整体呈现，避免过多分段，理想情况下每个主题部分应为一个连贯的段落
            3.控制整体字数，每个段落控制在150-200字左右，确保文书紧凑精炼
            4.第一段专业兴趣段落尤其需要注重逻辑性表述，避免堆砌过多素材
            5.增强句子之间的逻辑连接：
            - 确保每个新句子包含前一句子的关键词或概念
            - 使用指代词明确引用前文内容
            - 恰当使用过渡词和连接词
            - 建立清晰的因果关系，使用"因此"、"由此"、"正是"等词语明确前后句关系
            - 采用递进结构展示思想发展，从初始观察到深入思考，再到形成核心观点
            - 添加过渡句确保各点之间自然衔接，如"这种认识引导我..."、"通过这一探索..."- 确保每个段落形成完整的思想发展脉络，展现认知的深化过程   
            - 避免单纯并列不相关信息，而是通过逻辑词建立内在联系
            6.所有段落必须充分整合分析报告中的"个性化需求"和"增强策略"，同时保持简洁
            7.严格遵守报告中对每个段落的具体指导，不添加额外经历或合并不同经历
            8.强化内在逻辑推理，确保段落内容之间存在清晰的因果关系和思维发展脉络



            """,
            "material_simplifier_role": """
            该指令用于将个人陈述调查问卷中的零散信息转化为结构化的要点列表，以便于撰写留学申请材料。
            这一过程需要确保所有信息被正确归类，同时彻底移除任何学校和专业具体信息，以保持申请材料的通用性与适用性。
            留学申请中，个人陈述是展示申请者背景、经历、专业兴趣以及未来规划的关键材料，但原始调查问卷通常包含大量未经整理的信息，且可能包含过于具体的学校和专业信息，需要进行专业化的整理与归类。


            """,
            "material_simplifier_task": """
            1. 处理流程：
                - 仔细阅读提供的个人陈述调查问卷素材
                - 将素材中的信息按照统一格式提取
                - 删除学校和专业名称的同时保留项目实质内容
                - 按照个人陈述素材表的七大框架进行分类整理
                - 使用规定格式输出最终结果

            2. 关键要求：
                - 删除标识性信息但保留内容：
                    * 删除大学名称、缩写和别称，但保留在该校完成的项目、研究或经历的具体内容
                    * 删除实验室、研究中心的具体名称，但保留其研究方向和内容
                    * 删除特定学位项目名称和编号，但保留课程内容
                    * 删除教授、导师的姓名和头衔，但保留与其合作的项目内容
                
                - 保留课程和经历细节：
                    * 课程内容、项目描述、技能培养等细节必须完整保留
                    * 课程保留具体课程编号及课程名称
                    * 项目经历的技术细节、方法论、工具使用等信息必须保留
                    * 保留所有成果数据、获奖情况（移除具体学校名称）
                    * 即使项目是在特定学校完成的，也必须保留项目的全部实质内容
                
                - 信息分类必须精确无误：
                    * 每条信息必须且只能归入一个类别
                    * 严格遵循"七大框架"的分类标准
                    * 不允许创建新类别或合并现有类别
                    * 不允许同一信息跨类别重复出现
                
                - 经历要点格式要求：
                    * 研究、实习和实践经历必须按照七个子要点分行显示
                    * 如某些要素缺失，保持顺序不变并跳过该要素
                    * 项目内容描述必须包含项目所经历的完整步骤和流程
                    * 个人职责必须详细列出所有责任、遇到的困难及解决方案

            """,

            "material_simplifier_output":"""
            输出标题为"个人陈述素材整理报告"，仅包含以下七个部分：

            1. 专业兴趣塑造
            - 仅包含专业兴趣形成过程的要点列表
            - 每个要点以单个短横线"-"开头
            - 保留所有激发兴趣的细节经历和体验

            2. 学术基础展示
            - 仅包含课程学习、学术项目、教育背景的要点列表
            - 每个要点以单个短横线"-"开头
            - 保留课程内容、学习成果和技能培养的详细描述

            3. 研究经历深化
            - 每个研究经历作为一个主要要点，包含七个分行显示的子要点：
                - 项目名称：[内容]
                - 具体时间：[内容]
                - 扮演的角色：[内容]
                - 项目内容描述：[详细描述项目的全部步骤、背景、目标和实施过程]
                - 个人职责：[详细描述所有责任、遇到的困难及解决方案]
                - 取得的成果：[内容]
                - 经历感悟：[内容]

            4. 实习经历深化
            - 每个实习经历作为一个主要要点，包含七个分行显示的子要点：
                - 项目名称：[内容]
                - 具体时间：[内容]
                - 扮演的角色：[内容]
                - 项目内容描述：[详细描述项目的全部步骤、背景、目标和实施过程]
                - 个人职责：[详细描述所有责任、遇到的困难及解决方案]
                - 取得的成果：[内容]
                - 经历感悟：[内容]

            5. 实践经历补充
            - 每个实践经历作为一个主要要点，包含七个分行显示的子要点：
                - 项目名称：[内容]
                - 具体时间：[内容]
                - 扮演的角色：[内容]
                - 项目内容描述：[详细描述项目的全部步骤、背景、目标和实施过程]
                - 个人职责：[详细描述所有责任、遇到的困难及解决方案]
                - 取得的成果：[内容]
                - 经历感悟：[内容]

            6. 未来规划提升
            - 仅包含学习计划、职业规划、发展方向的要点列表
            - 每个要点以单个短横线"-"开头
            - 保留所有时间节点和具体规划细节

            7. 为何选择该专业和院校
            - 仅包含选择原因、国家优势的要点列表（不含具体学校信息）
            - 每个要点以单个短横线"-"开头
            - 保留对专业领域、研究方向的具体兴趣描述

            禁止添加任何非要求的标题、注释或总结。

            """
        }
        
        # 初始化 session_state 中的模板
        if 'templates' not in st.session_state:
            st.session_state.templates = self.default_templates.copy()

    def get_template(self, template_name: str) -> str:
        return st.session_state.templates.get(template_name, "")

    def update_template(self, template_name: str, new_content: str) -> None:
        st.session_state.templates[template_name] = new_content

    def reset_to_default(self):
        st.session_state.templates = self.default_templates.copy()

class TranscriptAnalyzer:
    def __init__(self, api_key: str, prompt_templates: PromptTemplates):
        self.prompt_templates = prompt_templates
        if 'templates' not in st.session_state:
            st.session_state.templates = self.prompt_templates.default_templates.copy()
            
        self.llm = ChatOpenAI(
            temperature=0.1,
            model=st.session_state.transcript_model,  # 使用session state中的模型
            api_key=api_key,
            base_url="https://openrouter.ai/api/v1",
            streaming=True
        )
        
        # 添加材料简化器LLM，使用成本较低的模型
        self.simplifier_llm = ChatOpenAI(
            temperature=0.1,
            model=st.session_state.simplifier_model,  # 使用session state中的简化器模型
            api_key=api_key,
            base_url="https://openrouter.ai/api/v1",
            streaming=True
        )
        self.setup_simplifier_chains()

    def extract_images_from_pdf(self, pdf_bytes):
        """从PDF中提取图像"""
        try:
            images = []
            pdf_document = fitz.open(stream=pdf_bytes, filetype="pdf")
            
            for page_num in range(len(pdf_document)):
                page = pdf_document[page_num]
                # 将页面直接转换为图像
                pix = page.get_pixmap(matrix=fitz.Matrix(2, 2))
                img_bytes = pix.tobytes("png")
                # 将图像编码为base64字符串
                img_base64 = base64.b64encode(img_bytes).decode('utf-8')
                images.append(img_base64)
            
            return images
        except Exception as e:
            logger.error(f"提取PDF图像时出错: {str(e)}")
            return []
    
    def analyze_transcript(self, pdf_bytes) -> Dict[str, Any]:
        try:
            if not hasattr(self, 'prompt_templates'):
                logger.error("prompt_templates not initialized")
                raise ValueError("Prompt templates not initialized properly")
            
            images = self.extract_images_from_pdf(pdf_bytes)
            if not images:
                return {
                    "status": "error",
                    "message": "无法从PDF中提取图像"
                }
            
            # 修改消息格式
            messages = [
                SystemMessage(content=self.prompt_templates.get_template('transcript_role')),
                HumanMessage(content=[  # 注意这里改成了列表
                    {
                        "type": "text",
                        "text": f"\n\n请分析这份成绩单，提取成绩信息，并以表格形式输出成绩信息。"
                    },
                    {
                        "type": "image_url",
                        "image_url": {
                            "url": f"data:image/png;base64,{images[0]}"
                        }
                    }
                ])
            ]
            
            # 创建一个队列用于流式输出
            message_queue = Queue()
            
            # 创建自定义回调处理器
            class QueueCallbackHandler(BaseCallbackHandler):
                def __init__(self, queue):
                    self.queue = queue
                    super().__init__()
                
                def on_llm_new_token(self, token: str, **kwargs) -> None:
                    self.queue.put(token)
            
            # 创建一个生成器函数，用于流式输出
            def token_generator():
                while True:
                    try:
                        token = message_queue.get(block=False)
                        yield token
                    except Empty:
                        if not thread.is_alive() and message_queue.empty():
                            break
                    time.sleep(0.01)
            
            # 在单独的线程中运行分析
            def run_analysis():
                try:
                    # 调用LLM进行分析
                    chain = LLMChain(llm=self.llm, prompt=ChatPromptTemplate.from_messages(messages))
                    result = chain.run(
                        {},
                        callbacks=[QueueCallbackHandler(message_queue)]
                    )
                    
                    message_queue.put("\n\n成绩单分析完成！")
                    thread.result = result
                    return result
                    
                except Exception as e:
                    message_queue.put(f"\n\n错误: {str(e)}")
                    logger.error(f"成绩单分析错误: {str(e)}")
                    thread.exception = e
                    raise e
            
            # 启动线程
            thread = Thread(target=run_analysis)
            thread.start()
            
            # 用于流式输出的容器
            output_container = st.empty()
            
            # 流式输出
            with output_container:
                full_response = st.write_stream(token_generator())
            
            # 等待线程完成
            thread.join()
            
            # 清空原容器并使用markdown重新渲染完整响应
            if full_response:
                output_container.empty()
                output_container.markdown(full_response)
            
            # 获取结果
            if hasattr(thread, "exception") and thread.exception:
                raise thread.exception
            
            logger.info("成绩单分析完成")
            
            return {
                "status": "success",
                "transcript_analysis": full_response
            }
                
        except Exception as e:
            logger.error(f"成绩单分析错误: {str(e)}")
            return {
                "status": "error",
                "message": str(e)
            }
    
    def setup_simplifier_chains(self):
            # 简化素材表 Chain
            simplifier_prompt = ChatPromptTemplate.from_messages([
                ("system", f"{self.prompt_templates.get_template('material_simplifier_role')}\n\n"
                        f"任务:\n{self.prompt_templates.get_template('material_simplifier_tesk')}\n\n"
                        f"请按照以下格式输出:\n{self.prompt_templates.get_template('material_simplifier_output')}"),
                ("human", "素材表document_content：\n{document_content}")
            ])
            
            self.simplifier_chain = LLMChain(
                llm=self.simplifier_llm,
                prompt=simplifier_prompt,
                output_key="simplifier_result",
                verbose=True
            )
    
    def simplify_materials(self, document_content: str) -> Dict[str, Any]:
        """简化素材表内容"""
        try:
            # 创建一个队列用于流式输出
            message_queue = Queue()
            
            # 创建自定义回调处理器，继承自 BaseCallbackHandler
            class QueueCallbackHandler(BaseCallbackHandler):
                def __init__(self, queue):
                    self.queue = queue
                    super().__init__()
                
                def on_llm_new_token(self, token: str, **kwargs) -> None:
                    self.queue.put(token)
            
            # 创建一个生成器函数，用于流式输出
            def token_generator():
                while True:
                    try:
                        token = message_queue.get(block=False)
                        yield token
                    except Empty:
                        if not thread.is_alive() and message_queue.empty():
                            break
                    time.sleep(0.01)
            
            # 在单独的线程中运行LLM
            def run_llm():
                try:
                    result = self.simplifier_chain(
                        {
                            "document_content": document_content
                        },
                        callbacks=[QueueCallbackHandler(message_queue)]
                    )
                    # 将结果存储在线程对象中
                    thread.result = result
                    message_queue.put("\n\n简化完成！")
                    return result
                except Exception as e:
                    message_queue.put(f"\n\n错误: {str(e)}")
                    logger.error(f"简化素材表时出错: {str(e)}")
                    thread.exception = e
                    raise e
            
            # 启动线程
            thread = Thread(target=run_llm)
            thread.start()
            with st.expander("简化后的素材表", expanded=True):
                # 创建流式输出容器
                output_container = st.empty()
                
                # 流式输出
                with output_container:
                    full_response = st.write_stream(token_generator())
                
                # 等待线程完成
                thread.join()
                
                # 清空原容器并使用markdown重新渲染完整响应
                if full_response:
                # 处理可能存在的markdown代码块标记
                    if full_response.startswith("```markdown"):
                        # 移除开头的```markdown和结尾的```
                        full_response = full_response.replace("```markdown", "", 1)
                        if full_response.endswith("```"):
                            full_response = full_response[:-3]
                    
                    output_container.empty()
                    new_container = st.container()
                    with new_container:
                        st.markdown(full_response)
                
                # 获取结果
                if hasattr(thread, "exception") and thread.exception:
                    raise thread.exception
                
                logger.info("simplifier_result completed successfully")
                
                # 从 full_response 中提取分析结果
                processed_response = full_response
                if processed_response.startswith("```markdown"):
                    # 移除开头的```markdown和结尾的```
                    processed_response = processed_response.replace("```markdown", "", 1)
                    if processed_response.endswith("```"):
                        processed_response = processed_response[:-3]
                # 从 full_response 中提取分析结果
                return {
                    "status": "success",
                    "simplifier_result": processed_response
                }
                    
        except Exception as e:
            logger.error(f"simplifier_result processing error: {str(e)}")
            return {
                "status": "error",
                "message": str(e)
            }

class BrainstormingAgent:
    def __init__(self, api_key: str, prompt_templates: PromptTemplates):

        self.content_llm = ChatOpenAI(
            temperature=0.1,
            model=st.session_state.content_model,  # 使用session state中的模型
            api_key=api_key,
            base_url="https://openrouter.ai/api/v1",
            streaming=True
        )
        self.prompt_templates = prompt_templates
        self.setup_chains()

    def setup_chains(self):        # 内容规划 Chain 
        creator_prompt = ChatPromptTemplate.from_messages([
            ("system", f"{self.prompt_templates.get_template('consultant_role2')}\n\n"
                      f"任务:\n{self.prompt_templates.get_template('consultant_task2')}\n\n"
                      f"请按照以下格式输出:\n{self.prompt_templates.get_template('output_format2')}"),
            ("human", "基于素材表document_content_simple：\n{document_content_simple}\n\n"
                     "成绩单transcript_analysis：\n{transcript_analysis}\n\n"
                     "申请方向school_plan：\n{school_plan}\n\n"
                     "定制需求custom_requirements：\n{custom_requirements}\n\n"
                     "请创建详细的内容规划。")
        ])
        
        self.creator_chain = LLMChain(
            llm=self.content_llm,
            prompt=creator_prompt,
            output_key="creator_output",
            verbose=True
        )

    def process_creator(self, document_content_simple: str, school_plan: str, transcript_analysis: str = "无成绩单", custom_requirements: str = "无定制需求") -> Dict[str, Any]:
        try:
            # 创建一个队列用于流式输出
            message_queue = Queue()
            
            # 创建自定义回调处理器，继承自 BaseCallbackHandler
            class QueueCallbackHandler(BaseCallbackHandler):
                def __init__(self, queue):
                    self.queue = queue
                    super().__init__()
                
                def on_llm_new_token(self, token: str, **kwargs) -> None:
                    self.queue.put(token)
            
            # 创建一个生成器函数，用于流式输出
            def token_generator():
                while True:
                    try:
                        token = message_queue.get(block=False)
                        yield token
                    except Empty:
                        if not thread.is_alive() and message_queue.empty():
                            break
                    time.sleep(0.01)
            
            # 在单独的线程中运行LLM
            def run_llm():
                try:
                    result = self.creator_chain(
                        {
                            "document_content_simple": document_content_simple,  # 添加文档内容
                            "school_plan": school_plan,
                            "transcript_analysis": transcript_analysis,
                            "custom_requirements": custom_requirements
                        },
                        callbacks=[QueueCallbackHandler(message_queue)]
                    )
                    # 将结果存储在队列中
                    message_queue.put("\n\n规划完成！")
                    return result
                except Exception as e:
                    message_queue.put(f"\n\n错误: {str(e)}")
                    logger.error(f"Creator processing error: {str(e)}")
                    raise e
            
            # 启动线程
            thread = Thread(target=run_llm)
            thread.start()
            
            # 创建流式输出容器
            output_container = st.empty()
            
            # 流式输出
            with output_container:
                full_response = st.write_stream(token_generator())
            
            # 等待线程完成
            thread.join()
            # 清空原容器并使用markdown重新渲染完整响应
            if full_response:
                # 处理可能存在的markdown代码块标记
                if full_response.startswith("```markdown"):
                    # 移除开头的```markdown和结尾的```
                    full_response = full_response.replace("```markdown", "", 1)
                    if full_response.endswith("```"):
                        full_response = full_response[:-3]
                
                output_container.empty()
                new_container = st.container()
                with new_container:
                    st.markdown(full_response)
            # 获取结果
            if hasattr(thread, "_exception") and thread._exception:
                raise thread._exception
            
            logger.info("Creator analysis completed successfully")
            processed_response = full_response
            if processed_response.startswith("```markdown"):
                # 移除开头的```markdown和结尾的```
                processed_response = processed_response.replace("```markdown", "", 1)
                if processed_response.endswith("```"):
                    processed_response = processed_response[:-3]

            return {
                "status": "success",
                "creator_output": processed_response
            }
                
        except Exception as e:
            logger.error(f"Creator processing error: {str(e)}")
            return {
                "status": "error",
                "message": str(e)
            }


def add_custom_css():
    st.markdown("""
    <style>
    /* 整体页面样式 */
    .main {
        padding: 2rem;
    }
    
    /* 标题样式 */
    h1, h2, h3 {
        color: #1e3a8a;
        font-weight: 600;
        margin-bottom: 1.5rem;
    }
    
    .page-title {
        text-align: center;
        font-size: 2.5rem;
        margin-bottom: 2rem;
        color: #1e3a8a;
        font-weight: bold;
        padding: 1rem;
        border-bottom: 3px solid #e5e7eb;
    }
    
    /* 文件上传区域样式 */
    .stFileUploader {
        margin-bottom: 2rem;
    }
    
    .stFileUploader > div > button {
        background-color: #f8fafc;
        color: #1e3a8a;
        border: 2px dashed #1e3a8a;
        border-radius: 8px;
        padding: 1rem;
        transition: all 0.3s ease;
    }
    
    .stFileUploader > div > button:hover {
        background-color: #f0f7ff;
        border-color: #2563eb;
    }
    
    /* 按钮样式 */
    .stButton > button {
        background-color: #1e3a8a;
        color: white;
        border-radius: 8px;
        padding: 0.75rem 1.5rem;
        font-weight: 500;
        border: none;
        width: 100%;
        transition: all 0.3s ease;
        box-shadow: 0 2px 4px rgba(0,0,0,0.1);
    }
    
    .stButton > button:hover {
        background-color: #2563eb;
        transform: translateY(-1px);
        box-shadow: 0 4px 6px rgba(0,0,0,0.1);
    }
    
    .stButton > button:disabled {
        background-color: #94a3b8;
        cursor: not-allowed;
    }
    
    /* 文本区域样式 */
    .stTextArea > div > div > textarea {
        border-radius: 8px;
        border: 1px solid #e2e8f0;
        padding: 0.75rem;
        font-size: 1rem;
        transition: all 0.3s ease;
    }
    
    .stTextArea > div > div > textarea:focus {
        border-color: #2563eb;
        box-shadow: 0 0 0 3px rgba(37,99,235,0.1);
    }
    
    /* 分析结果区域样式 */
    .analysis-container {
        background-color: white;
        border-radius: 12px;
        padding: 1.5rem;
        margin: 1rem 0;
        box-shadow: 0 4px 6px rgba(0,0,0,0.05);
        border: 1px solid #e5e7eb;
    }
    
    /* 成功消息样式 */
    .stSuccess {
        background-color: #ecfdf5;
        color: #065f46;
        padding: 1rem;
        border-radius: 8px;
        border-left: 4px solid #059669;
        margin: 1rem 0;
    }
    
    /* 错误消息样式 */
    .stError {
        background-color: #fef2f2;
        color: #991b1b;
        padding: 1rem;
        border-radius: 8px;
        border-left: 4px solid #dc2626;
        margin: 1rem 0;
    }
    
    /* 模型信息样式 */
    .model-info {
        background-color: #f0f7ff;
        padding: 0.75rem 1rem;
        border-radius: 8px;
        margin: 0.5rem 0;
        display: inline-block;
        font-size: 0.9rem;
        border: 1px solid #bfdbfe;
    }
    
    /* 双列布局样式 */
    .dual-column {
        display: flex;
        gap: 2rem;
        margin: 1rem 0;
    }
    
    .column {
        flex: 1;
        background-color: #f8fafc;
        padding: 1.5rem;
        border-radius: 12px;
        border: 1px solid #e5e7eb;
    }
    
    /* 分隔线样式 */
    hr {
        margin: 2rem 0;
        border: 0;
        height: 1px;
        background: linear-gradient(to right, transparent, #e5e7eb, transparent);
    }
    
    /* 标签页样式 */
    .stTabs {
        background-color: white;
        border-radius: 12px;
        padding: 1rem;
        box-shadow: 0 1px 3px rgba(0,0,0,0.1);
    }
    
    .stTab {
        padding: 1rem;
    }
    
    /* 展开器样式 */
    .streamlit-expanderHeader {
        background-color: #f8fafc;
        border-radius: 8px;
        padding: 0.75rem;
        font-weight: 500;
        color: #1e3a8a;
        border: 1px solid #e5e7eb;
    }
    
    .streamlit-expanderContent {
        background-color: white;
        border-radius: 0 0 8px 8px;
        padding: 1rem;
        border: 1px solid #e5e7eb;
        border-top: none;
    }
    
    /* 加载动画样式 */
    .stSpinner > div {
        border-color: #2563eb transparent transparent transparent;
    }
    
    /* 文档分析区域样式 */
    .doc-analysis-area {
        background-color: #ffffff;
        border-radius: 12px;
        padding: 1.5rem;
        margin: 1rem 0;
        box-shadow: 0 2px 4px rgba(0,0,0,0.05);
        border: 1px solid #e5e7eb;
    }
    
    .doc-analysis-area h3 {
        color: #1e3a8a;
        font-size: 1.25rem;
        margin-bottom: 1rem;
        padding-bottom: 0.5rem;
        border-bottom: 2px solid #e5e7eb;
    }
    
    /* 调整列宽度 */
    .column-adjust {
        padding: 0 1rem;
    }
    </style>
    """, unsafe_allow_html=True)


def process_document_with_markitdown(content):
    """使用Markitdown处理文档内容"""
    try:
        # 配置Markitdown
        return markitdown.convert(content, extensions=[
            'extra',      # 包含tables, attr_list, footnotes等
            'codehilite', # 代码高亮
            'toc',        # 目录生成
            'smarty',     # 智能标点
            'sane_lists'  # 更好的列表处理
        ], extension_configs={
            'codehilite': {
                'linenums': False,
                'css_class': 'highlight'
            },
            'toc': {
                'permalink': True,
                'title': '目录'
            }
        })
    except Exception as e:
        logger.error(f"使用Markitdown处理内容时出错: {str(e)}")
        return content  # 出错时返回原始内容

def read_docx(file_bytes):
    """读取 Word 文档内容，包括表格，使用markitdown处理后返回"""
    try:
        doc = Document(io.BytesIO(file_bytes))
        content_set = set()  # 用于存储已处理的内容，避免重复
        full_text = []
        
        # 读取普通段落
        for paragraph in doc.paragraphs:
            text = paragraph.text.strip()
            if text and text not in content_set:  # 只添加非空且未重复的内容
                content_set.add(text)
                full_text.append(text)
        
        # 读取表格内容
        for table in doc.tables:
            md_table = convert_table_to_markdown(table)
            if md_table:
                full_text.append(md_table)
        
        # 使用换行符连接所有文本，创建完整的Markdown文本
        raw_markdown = "\n\n".join(full_text)
        
        # 使用markitdown处理提取的内容
        processed_content = process_document_with_markitdown(raw_markdown)
        
        logger.info(f"成功读取文档内容，使用Markitdown处理完成，包含 {len(doc.tables)} 个表格")
        
        # 返回原始提取文本和处理后的HTML内容
        return {
            'raw_content': raw_markdown,
            'html_content': processed_content
        }
    except Exception as e:
        logger.error(f"读取 Word 文档时出错: {str(e)}")
        return None

def convert_table_to_markdown(table):
    """将Word表格转换为Markdown表格格式"""
    if not table.rows:
        return ""
    
    md_table = []
    header = []
    
    # 获取表头
    for cell in table.rows[0].cells:
        header.append(cell.text.strip())
    
    if not header:
        return ""
    
    # 添加表头行
    md_table.append("| " + " | ".join(header) + " |")
    
    # 添加分隔行
    md_table.append("| " + " | ".join(["---" for _ in header]) + " |")
    
    # 添加数据行
    for row in table.rows[1:]:
        cells = [cell.text.strip() for cell in row.cells]
        if any(cells):  # 确保行不是空的
            # 确保单元格数量与表头一致
            while len(cells) < len(header):
                cells.append("")
            cells = cells[:len(header)]  # 截断多余的单元格
            md_table.append("| " + " | ".join(cells) + " |")
    
    return "\n".join(md_table)

def initialize_session_state():
    if 'templates' not in st.session_state:
        prompt_templates = PromptTemplates()
        st.session_state.templates = prompt_templates.default_templates.copy()
    
    if 'prompt_templates' not in st.session_state:
        st.session_state.prompt_templates = PromptTemplates()

def main():
    initialize_session_state()
    
    langsmith_api_key = st.secrets["LANGCHAIN_API_KEY"]
    os.environ["LANGCHAIN_TRACING_V2"] = "true"
    os.environ["LANGCHAIN_API_KEY"] = langsmith_api_key
    os.environ["LANGCHAIN_PROJECT"] = "初稿脑暴平台"
    st.set_page_config(page_title="初稿脑暴助理平台", layout="wide")
    add_custom_css()
    st.markdown("<h1 class='page-title'>初稿脑暴助理</h1>", unsafe_allow_html=True)
    
    # 添加模型选择
    col1, col2, col3 = st.columns(3)
    with col1:
        if 'transcript_model' not in st.session_state:
            st.session_state.transcript_model = st.secrets["TRANSCRIPT_MODEL"]
        
        transcript_model = st.selectbox(
            "选择图像分析模型",
            ["google/gemma-3-27b-it:free","google/gemini-2.0-flash-001","google/gemini-2.5-pro-exp-03-25:free", "qwen/qwen2.5-vl-32b-instruct:free","anthropic/claude-3.7-sonnet"],
            index=["google/gemma-3-27b-it:free","google/gemini-2.0-flash-001","google/gemini-2.5-pro-exp-03-25:free", "qwen/qwen2.5-vl-32b-instruct:free","anthropic/claude-3.7-sonnet"].index(st.session_state.transcript_model)
        )
        if transcript_model != st.session_state.transcript_model:
            st.session_state.transcript_model = transcript_model
            st.rerun()
    
    with col2:
        if 'simplifier_model' not in st.session_state:
            st.session_state.simplifier_model = st.secrets["SIMPLIFIER_MODEL"]
        
        simplifier_model = st.selectbox(
            "选择简化模型",
            ["qwen/qwq-32b:free","qwen/qwq-32b","google/gemini-2.5-pro-exp-03-25:free", "deepseek/deepseek-chat-v3-0324:free", "deepseek/deepseek-r1:free","deepseek/deepseek-r1","anthropic/claude-3.7-sonnet"],
            index=["qwen/qwq-32b:free","qwen/qwq-32b","google/gemini-2.5-pro-exp-03-25:free", "deepseek/deepseek-chat-v3-0324:free", "deepseek/deepseek-r1:free","deepseek/deepseek-r1","anthropic/claude-3.7-sonnet"].index(st.session_state.simplifier_model)
        )
        if simplifier_model != st.session_state.simplifier_model:
            st.session_state.simplifier_model = simplifier_model
            st.rerun()
            
    with col3:
        if 'content_model' not in st.session_state:
            st.session_state.content_model = st.secrets["CONTENT_MODEL"]
        content_model = st.selectbox(
            "选择内容规划模型",
            ["qwen/qwq-32b:free","qwen/qwq-32b","google/gemini-2.5-pro-exp-03-25:free", "deepseek/deepseek-chat-v3-0324:free", "deepseek/deepseek-r1:free","deepseek/deepseek-r1","anthropic/claude-3.7-sonnet"],
            index=["qwen/qwq-32b:free","qwen/qwq-32b","google/gemini-2.5-pro-exp-03-25:free", "deepseek/deepseek-chat-v3-0324:free", "deepseek/deepseek-r1:free","deepseek/deepseek-r1","anthropic/claude-3.7-sonnet"].index(st.session_state.content_model)
        )
        if content_model != st.session_state.content_model:
            st.session_state.content_model = content_model
            st.rerun()
            
    # 添加单个模型信息展开框
    with st.expander("查看模型详细信息", expanded=False):
        st.markdown("""
        ### 图像分析模型
        
        **google/gemma-3-27b-it:free**
        - 参数量：27B
        - Created Mar 12, 2025 | 96,000 context | $0/M  input tokens | $0\/M output tokens
        
        **google/gemini-2.0-flash-001**
        - Created Feb 5, 2025 |  1,000,000 context |  $0.1/M  input tokens  | $0.4\/M output tokens |  $0.0258\/K input imgs
        
        **google/gemini-2.5-pro-exp-03-25:free**
        - 最新版本的Gemini模型
        - Created Mar 25, 2025 | 1,000,000 context | $0/M  input tokens | $0\/M output tokens
        
        **qwen/qwen2.5-vl-32b-instruct:free**
        - 参数量：32B
        - Created Mar 24, 2025 | 8,192 context | $0/M  input tokens | $0\/M output tokens
                    
        **anthropic/claude-3.7-sonnet**
        - 参数量：32B
        - Created Feb 24, 2025 | 200,000 context | $3/M  input tokens | $15\/M output tokens | $4.8\/K input imgs
        
        ### 文本分析模型
        
        **qwen/qwq-32b:free**
        - 参数量：32B
        - Created Mar 5, 2025 | 40,000 context | $0/M input tokens | $0\/M output tokens
        
        **qwen/qwq-32b**
        - 完整版本，性能更优
        - Created Mar 5, 2025 | 131,072 context | $0.12/M input tokens  | $0.18\/M output tokens
        
        **google/gemini-2.5-pro-exp-03-25:free**
        - 最新版本的Gemini模型
        - Created Mar 25, 2025 | 1,000,000 context | $0/M input tokens | $0\/M output tokens
        
        **deepseek/deepseek-chat-v3-0324:free**
        - 最新的v3版本
        - Created Mar 24, 2025 | 131,072 context | $0/M input tokens | $0\/M output tokens
        
        **deepseek/deepseek-r1:free & deepseek/deepseek-r1**
        - Created Jan 20, 2025 | 163,840 context | $0/M input tokens | $0\/M output tokens
                    
        **anthropic/claude-3.7-sonnet**
        - 参数量：32B
        - Created Feb 24, 2025 | 200,000 context | $3/M input tokens | $15\/M output tokens | $4.8\/K input imgs
        
        """)

    # 确保在任何操作之前初始化 PromptTemplates
    if 'templates' not in st.session_state:
        prompt_templates = PromptTemplates()
        st.session_state.templates = prompt_templates.default_templates.copy()
    
    if 'prompt_templates' not in st.session_state:
        st.session_state.prompt_templates = PromptTemplates()
    
    tab1, tab2 = st.tabs(["初稿脑暴助理", "提示词设置"])
    
    # 初始化会话状态变量
    if 'document_content' not in st.session_state:
        st.session_state.document_content = None
    if 'transcript_file' not in st.session_state:
        st.session_state.transcript_file = None
    if 'transcript_analysis_done' not in st.session_state:
        st.session_state.transcript_analysis_done = False
    if 'transcript_analysis_result' not in st.session_state:
        st.session_state.transcript_analysis_result = None
    if 'strategist_analysis_done' not in st.session_state:
        st.session_state.strategist_analysis_done = False
    if 'strategist_analysis_result' not in st.session_state:
        st.session_state.strategist_analysis_result = None
    if 'creator_analysis_done' not in st.session_state:
        st.session_state.creator_analysis_done = False
    if 'creator_analysis_result' not in st.session_state:
        st.session_state.creator_analysis_result = None
    if 'show_transcript_analysis' not in st.session_state:
        st.session_state.show_transcript_analysis = False
    if 'show_strategist_analysis' not in st.session_state:
        st.session_state.show_strategist_analysis = False
    if 'show_creator_analysis' not in st.session_state:
        st.session_state.show_creator_analysis = False
    if 'show_simplifier_analysis' not in st.session_state:
        st.session_state.show_simplifier_analysis = False
    if 'simplifier_analysis_done' not in st.session_state:
        st.session_state.simplifier_analysis_done = False
    if 'simplifier_result' not in st.session_state:
        st.session_state.simplifier_result = None
    
    with tab1:
        transcript_file = st.file_uploader("上传成绩单（可选）", type=['pdf'])
        # 自动检查文件状态并清除相关内存
        if not transcript_file:
            st.session_state.transcript_file = None
            st.session_state.transcript_analysis_done = False
            st.session_state.transcript_analysis_result = None
            st.session_state.show_transcript_analysis = False
        
        # 添加分析成绩单按钮
        if st.button("分析成绩单", key="analyze_transcript", use_container_width=True):
            if transcript_file is not None:
                st.session_state.transcript_file = transcript_file
                st.session_state.show_transcript_analysis = True
                st.session_state.transcript_analysis_done = False
                st.rerun()
        
        # 修改文件上传部分，移除多文件支持
        
        
        uploaded_file = st.file_uploader("上传初稿文档", type=['docx'])  # 改为单文件上传
        
        # 处理上传的文件
        if uploaded_file:
            raw_content = read_docx(uploaded_file.read())
            if raw_content:
                # 使用Markitdown处理内容
                html_content = process_document_with_markitdown(raw_content['raw_content'])
                
                # 保存原始内容用于后续分析
                st.session_state.document_content = raw_content['raw_content']
                
                # 保存处理后的HTML内容，如果需要展示
                st.session_state.document_html = html_content
                
                st.success(f"文件上传并处理成功！")
                
                # 显示处理结果
                with st.expander("查看Markitdown处理结果", expanded=False):
                    st.markdown(html_content, unsafe_allow_html=True)
                    
                with st.expander("查看原始提取内容", expanded=False):
                    st.write(raw_content['raw_content'])
            else:
                st.error("无法读取文件，请检查格式是否正确。")
        
        if st.button("简化素材表", key="simplify_materials", use_container_width=True):
            if st.session_state.document_content:
                st.session_state.show_simplifier_analysis = True
                st.session_state.simplifier_analysis_done = False
                st.rerun()
        if st.session_state.show_simplifier_analysis:
            with st.container():
                st.markdown("---")
                st.subheader("📊 简化后的素材表")
                
                if not st.session_state.simplifier_analysis_done:
                    try:
                        # 确保 prompt_templates 存在
                        if 'prompt_templates' not in st.session_state:
                            st.session_state.prompt_templates = PromptTemplates()
                        
                        transcript_analyzer = TranscriptAnalyzer(
                            api_key=st.secrets["OPENROUTER_API_KEY"],  # 使用OpenRouter API密钥
                            prompt_templates=st.session_state.prompt_templates
                        )
                        
                        with st.spinner("正在简化素材表..."):
                            # 处理成绩单分析
                            result = transcript_analyzer.simplify_materials(
                                st.session_state.document_content
                            )
                            
                            if result["status"] == "success":
                                # 保存成绩单分析结果到 session_state
                                st.session_state.simplifier_result = result["simplifier_result"]
                                st.session_state.simplifier_analysis_done = True
                                st.success("✅ 简化素材表完成！")
                            else:
                                st.error(f"简化素材表出错: {result['message']}")
                    
                    except Exception as e:
                        st.error(f"处理过程中出错: {str(e)}")
                else:
                    with st.expander("查看简化后的素材表", expanded=False):
                        # 如果已经完成，直接显示结果
                        st.markdown(st.session_state.simplifier_result)
                        st.success("✅ 简化素材表完成！")
        # 添加选校方案输入框
        school_plan = st.text_area(
            "选校方案",
            value="暂未选校",
            height=100,
            help="请输入已确定的选校方案，包括学校和专业信息"
        )
        
        # 添加自定义需求输入框
        custom_requirements = st.text_area(
            "定制需求（可选）",
            value="无定制需求",
            height=100,
            help="请输入特殊的定制需求，如果没有可以保持默认值"
        )
        
        # 修改按钮区域，只保留单文件模式
        button_col1, button_col2 = st.columns(2)
        
        with button_col1:
            continue_button = st.button(
                "内容规划", 
                key="continue_to_creator", 
                use_container_width=True
            )
            
            if continue_button:
                st.session_state.show_creator_analysis = True
                st.session_state.creator_analysis_done = False
                st.rerun()
        with button_col2:
            if st.button("清除所有分析", key="clear_analysis", use_container_width=True):
                # 清除所有分析相关的session状态
                st.session_state.document_content = None
                st.session_state.strategist_analysis_done = False
                st.session_state.creator_analysis_done = False
                st.session_state.show_strategist_analysis = False
                st.session_state.show_creator_analysis = False
                st.session_state.show_simplifier_analysis = False
                st.session_state.simplifier_analysis_done = False
                st.success("✅ 所有分析结果已清除！")
                st.rerun()
        # 修改结果显示区域，只保留单文档逻辑
        results_container = st.container()
        
        # 显示成绩单分析（保持不变）
        if st.session_state.show_transcript_analysis:
            with results_container:
                st.markdown("---")
                st.subheader("📊 成绩单分析")
                
                if not st.session_state.transcript_analysis_done:
                    try:
                        # 确保 prompt_templates 存在
                        if 'prompt_templates' not in st.session_state:
                            st.session_state.prompt_templates = PromptTemplates()
                        
                        transcript_analyzer = TranscriptAnalyzer(
                            api_key=st.secrets["OPENROUTER_API_KEY"],  # 使用OpenRouter API密钥
                            prompt_templates=st.session_state.prompt_templates
                        )
                        
                        with st.spinner("正在分析成绩单..."):
                            # 处理成绩单分析
                            result = transcript_analyzer.analyze_transcript(
                                st.session_state.transcript_file
                            )
                            
                            if result["status"] == "success":
                                # 保存成绩单分析结果到 session_state
                                st.session_state.transcript_analysis_result = result["transcript_analysis"]
                                st.session_state.transcript_analysis_done = True
                                st.success("✅ 成绩单分析完成！")
                            else:
                                st.error(f"成绩单分析出错: {result['message']}")
                    
                    except Exception as e:
                        st.error(f"处理过程中出错: {str(e)}")
                else:
                    # 如果已经完成，直接显示结果
                    st.markdown(st.session_state.transcript_analysis_result)
                    st.success("✅ 成绩单分析完成！")
        
        # 修改内容规划显示，只保留单文档逻辑
        if st.session_state.show_creator_analysis:
            with results_container:
                st.markdown("---")
                st.subheader("📝 内容规划")
                
                if not st.session_state.creator_analysis_done:
                    try:
                        agent = BrainstormingAgent(
                            api_key=st.secrets["OPENROUTER_API_KEY"],
                            prompt_templates=st.session_state.prompt_templates
                        )
                        document_content_simple = ""
                        if st.session_state.simplifier_result == None:
                            document_content_simple = st.session_state.document_content
                            st.write("使用原始素材表进行分析")
                        else:
                            document_content_simple = st.session_state.simplifier_result
                            st.write("使用简化后素材表进行分析")
                        with st.spinner("正在规划内容..."):
                            creator_result = agent.process_creator(
                                document_content_simple = document_content_simple,
                                school_plan=school_plan,
                                transcript_analysis=st.session_state.transcript_analysis_result,
                                custom_requirements=custom_requirements
                            )
                            
                            if creator_result["status"] == "success":
                                st.session_state.creator_analysis_result = creator_result["creator_output"]
                                st.session_state.creator_analysis_done = True
                                st.success("✅ 内容规划完成！")
                            else:
                                st.error(f"内容规划出错: {creator_result['message']}")
                    except Exception as e:
                        st.error(f"处理过程中出错: {str(e)}")
                else:
                    # 使用markdown方法并明确指定unsafe_allow_html参数
                    st.markdown(st.session_state.creator_analysis_result, unsafe_allow_html=True)
                    st.success("✅ 内容规划完成！")
            # 修改模型信息显示
        st.markdown(f"<div class='model-info'>🤖 图像分析当前使用模型: <b>{st.session_state.transcript_model}</b></div>", unsafe_allow_html=True)
        st.markdown(f"<div class='model-info'>🤖 内容规划当前使用模型: <b>{st.session_state.content_model}</b></div>", unsafe_allow_html=True)
        
    
    with tab2:
        st.title("提示词设置")
        
        prompt_templates = st.session_state.prompt_templates
        
        #素材表简化提示词
        st.subheader("素材表简化")
        material_simplifier_role = st.text_area(
            "角色设定",
            value=prompt_templates.get_template('material_simplifier_role'),
            height=200,
            key="material_simplifier_role"
        )
        material_simplifier_task = st.text_area(
            "任务说明",
            value=prompt_templates.get_template('material_simplifier_task'),
            height=200,
            key="material_simplifier_task"
        )
        material_simplifier_output = st.text_area(
            "输出格式",
            value=prompt_templates.get_template('material_simplifier_output'),
            height=200,
            key="material_simplifier_output"
        )
        
        # Agent 2 设置
        st.subheader("Agent 2 - 内容创作师")
        consultant_role2 = st.text_area(
            "角色设定",
            value=prompt_templates.get_template('consultant_role2'),
            height=200,
            key="consultant_role2"
        )

        consultant_task2 = st.text_area(
            "任务说明",
            value=prompt_templates.get_template('consultant_task2'),
            height=200,
            key="consultant_task2"
        )

        output_format2 = st.text_area(
            "输出格式",
            value=prompt_templates.get_template('output_format2'),
            height=200,
            key="output_format2"
        )
        
        col1, col2 = st.columns(2)
        with col1:
            if st.button("更新提示词", key="update_prompts"):
                prompt_templates.update_template('material_simplifier_role', material_simplifier_role)
                prompt_templates.update_template('material_simplifier_task', material_simplifier_task)
                prompt_templates.update_template('material_simplifier_output', material_simplifier_output)
                prompt_templates.update_template('consultant_role2', consultant_role2)
                prompt_templates.update_template('output_format2', output_format2)
                prompt_templates.update_template('consultant_task2', consultant_task2)
                st.success("✅ 提示词已更新！")
        
        with col2:
            if st.button("重置为默认提示词", key="reset_prompts"):
                prompt_templates.reset_to_default()
                st.rerun()

if __name__ == "__main__":
    main()