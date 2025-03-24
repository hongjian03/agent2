import streamlit as st
import json
from langchain_openai import ChatOpenAI
from langchain.prompts import ChatPromptTemplate
from langchain.schema import HumanMessage, SystemMessage
from langchain.callbacks.streaming_stdout import StreamingStdOutCallbackHandler
from langchain.callbacks.streamlit import StreamlitCallbackHandler
from langchain.chains import SequentialChain, LLMChain
import os
from typing import Dict, Any, List
import logging
import sys
from docx import Document
import io
import base64
from PyPDF2 import PdfReader
from PIL import Image
import fitz  # PyMuPDF
# 配置日志记录
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(levelname)s - %(message)s',
    handlers=[
        logging.StreamHandler(sys.stdout)
    ]
)
from queue import Queue
from threading import Thread
import time
from queue import Empty
logger = logging.getLogger(__name__)
from langchain.callbacks.base import BaseCallbackHandler

# 记录程序启动
logger.info("程序开始运行")

# 只在第一次运行时替换 sqlite3
if 'sqlite_setup_done' not in st.session_state:
    try:
        logger.info("尝试设置 SQLite")
        __import__('pysqlite3')
        sys.modules['sqlite3'] = sys.modules.pop('pysqlite3')
        st.session_state.sqlite_setup_done = True
        logger.info("SQLite 设置成功")
    except Exception as e:
        logger.error(f"SQLite 设置错误: {str(e)}")
        st.session_state.sqlite_setup_done = True


class PromptTemplates:
    def __init__(self):
        # 定义示例数据作为字符串
        self.default_templates = {
            'transcript_role': """
            # 角色
            你是专业的成绩单分析师，擅长从成绩单中提取关键信息并以表格形式展示成绩。
            """,
            
            'transcript_task': """
            分析学生的成绩单，提取以下信息：
            1. 学生的GPA和成绩分布情况
            2. 主要课程的成绩表现
            3. 学术优势和劣势
            4. 成绩趋势（是否有进步或下滑）
            5. 与申请专业相关课程的表现
            """,
            
            'transcript_output': """
            成绩单分析:
                GPA和总体表现: [GPA和总体成绩分布]
                主要课程成绩: [列出主要课程及成绩]
                学术优势: [分析学生的学术优势]
                学术劣势: [分析学生的学术劣势]
                成绩趋势: [分析成绩的变化趋势]
                与申请专业相关性: [分析与申请专业相关课程的表现]
            """,
            
            'consultant_role1': """
            # 角色
            我是一位拥有十年经验的顶尖大学招生文书评估专家。我擅长分析申请专业关联性，帮助学生找到最有竞争力的申请文书策略，减少重复工作并提高申请效率。

            """,
            
            'output_format1': """
            必须完整分析并输出所有识别出的专业方向（最多2个），每个方向的分析必须包含以下内容：
            
            个人陈述初稿分析报告
            
            一、专业方向分析
            专业方向1：[名称]
            所属院校及专业清单：
                ● 招生倾向分析：
                ●[招生倾向点1]
                ●[招生倾向点2]
                ●[招生倾向点3]
                ●[...]
            
            专业方向2：[名称] (如识别出第二个专业方向，必须完整分析)
            所属院校及专业清单：
                ● 招生倾向分析：
                ●[招生倾向点1]
                ●[招生倾向点2]
                ●[招生倾向点3]
                ●[...]
            
            二、段落素材策略与增强指南
            必须为每个专业方向分别提供完整的段落策略分析：
            
            专业方向1
            1. 专业兴趣段落
            ● 选择素材：[选择的具体素材]
            ● 现有基础：[简述现有素材]
            ● 个性化需求：[根据用户提供的个性化需求分析判断，如无则标注"无特殊要求"]
            ● 增强策略：
            - A. 确定写作角度（必须选择一个主要角度）：
            - 基于过去经历谈专业理解：选择非核心经历，避免与后续科研/实习段落重复使用重要素材
            - 基于时事新闻谈专业理解：关联当前专业领域热点话题
            - 基于科研成果谈专业前景：分析专业未来发展方向和个人兴趣点
            - B. [具体增强方法B]
            - C. [具体增强方法C]
            - D. [具体增强方法D]
            
            2. 学术基础展示
            ● 选择素材：[选择的具体素材]
            ● 现有基础：[简述现有素材]
            ● 个性化需求：[根据用户提供的个性化需求分析判断，如无则标注"无特殊要求"]
            ● 增强策略：
            - A. [具体增强方法A]
            - B. [具体增强方法B]
            - C. [具体增强方法C]
            - D. [具体增强方法D]
            
            3. 科研经历深化
            ● 选择素材：[选择的具体素材]
            ● 现有基础：[简述现有素材]
            ● 个性化需求：[根据用户提供的个性化需求分析判断，如无则标注"无特殊要求"]
            ● 增强策略：
            - A.明确项目背景：
            - [详细说明项目的具体情况]
            - [解释研究主题选择原因]
            - B.深化研究过程：
            - [描述研究方法学]
            - [增加研究中遇到的挑战及解决方法]
            - [添加与导师/专家交流内容]
            - C.强化研究成果：
            - [明确研究发现和建议]
            - [将研究与专业方向关联]
            - [分析研究对专业理解的影响]
            - D.联系申请专业：
            - [分析研究如何启发对该专业的兴趣]
            - [强调研究方法与专业的相关性]
            
            4. 实习经历深化
            ● 选择素材：[必须只选择一个最相关的实习经历]
            ● 现有基础：[简述现有素材]
            ● 个性化需求：[根据用户提供的个性化需求分析判断，如无则标注"无特殊要求"]
            ● 增强策略：
            - A.明确项目结构：
            - [详细说明项目的具体目标和框架]
            - [描述项目如何与专业方向相关]
            - B.强化个人贡献：
            - [具体阐述职责，使用ACTION-RESULT模式]
            - [量化成就]
            - [详细描述参与的具体案例]
            - C.深化专业见解：
            - [分析工作中面临的专业挑战]
            - [反思专业相关问题]
            - [展示对行业理解]
            - D.与申请专业联系：
            - [说明经历如何引导对专业的兴趣]
            - [解释为何需要该专业的系统知识]
            - [表明经历如何确认了专业需求]
            
            5. 未来规划提升
            ● 选择素材：[选择的具体素材]
            ● 现有基础：[简述现有素材]
            ● 个性化需求：[根据用户提供的个性化需求分析判断，如无则标注"无特殊要求"]
            ● 增强策略：
            - A.明确短期目标：
            - [详细说明希望获得的具体技能和知识]
            - [计划获取的特定认证或参与的项目]
            - [希望建立的学术和专业网络]
            - B.设计中期职业路径：
            - [毕业后1-3年的具体职业目标]
            - [希望解决的专业相关挑战]
            - [工作地理区域意向]
            - C.制定长期愿景：
            - [5-10年职业发展愿景]
            - [希望在专业领域的具体贡献]
            - [与申请专业训练的关联]
            - D.增加可行性：
            - [如何利用现有资源和网络]
            - [毕业后的具体目标组织或项目]
            - [展示对行业趋势的了解]
            
            专业方向2 (如识别出第二个专业方向，必须按照相同格式完整分析)
            [与专业方向1相同的段落结构]
            
            重要说明：
            1.必须严格按照上述格式输出分析报告
            2.如果用户提供明确的申请专业方向，则必须直接使用且不进行任何修改或补充，绝不从素材表中重新解读或与素材表中的信息结合；如提供选校规划，则需进行分析和归类
            3.当用户明确提供申请方向时，不得基于素材表中的任何信息替换或修改该方向
            4.如果识别出两个专业方向，必须为两个方向分别提供完整的分析
            5.科研经历、实习经历必须只选择一个最相关的经历，不得选择多个
            6.所有分析必须基于学生提供的真实材料
            7.所有"个性化需求"必须根据用户实际提供的需求进行分析判断，如无特殊要求则标注"无特殊要求"
            8.对于"增加经历"类需求，必须首先检查是否有未使用的相关素材，只有在没有可用素材时才创建新内容
            9.增强策略必须按照固定格式分点展示，确保清晰易读
            10.根据用户的实际需求灵活调整格式，确保所有增加的经历都被有效整合到相应的段落中，每个经历都应有完整的分析结构

            """,
            
            'consultant_task1': """
            基本任务：
            1.仔细阅读和充分理解用户上传的素材表document_content、输入的个性化需求custom_requirements、分析用户上传的成绩单transcript_analysis、分析用户提供的申请专业方向或详细的选校规划school_plan
            2.准确识别用户指定的申请方向：
                - 如用户提供明确的申请方向名称（如"公共政策"），则必须严格以用户提供的名称为唯一分析对象，完全忽略素材表中的任何专业信息，不得结合或修改用户指定的方向名称
                - 如用户提供详细选校规划，则需分析各院校专业的共性与差异，归类为1-2个专业大类方向
                - 禁止根据用户提供的申请专业方向及心仪院校确认申请方向
                - 禁止将用户提供的素材表中的任何申请院校信息内容纳入选校规划的分析
            3.为每个识别出的专业大类设计一套完整的基础文书策略，包含专业兴趣塑造、学术基础展示、科研经历深化、实习经历深化、未来规划提升五个必要段落
            4.评估现有素材对各专业大类的适用性，提出具体可执行的优化建议

            重要规则：
            1.禁止提及任何具体学校名称，即使素材表中提及了相关学校和专业的内容，也要忽略学术基础展示部分不要讨论成绩
            2.学术基础展示部分不要讨论成绩，因为招生官可直接从成绩单查看
            3.当要求"增加经历"时，应遵循以下原则：
                a. 如果用户需求是增加新的经历段落，则保持原有段落不变，并添加一个新的完整段落，使用"XX经历深化1"、"XX经历深化2"等标题格式
                b. 如果用户需求是在现有段落中增加多个经历，则在同一段落内使用子标题"XX经历1"、"XX经历2"等进行区分
                c. 增加的经历段落或子段落必须保持与其他段落相同的结构和深度
            4.此规则适用于所有类型的经历（科研、实习、课外活动等），每个经历段落或子段落都必须有完整的分析结构

            分析流程：
            第一步：专业方向分析
            1.判断用户输入类型：
                - 如果是明确的申请专业方向名称，必须以用户明确提供的申请方向为唯一分析对象，完全忽略素材表中的专业方向信息，不得对用户提供的方向名称进行任何修改或补充
                - 如果用户提供选校规划（包含多个院校及专业名称），严格执行以下分析：
                    a. 分析表格中列出的所有院校专业名称
                    b. 将专业按学科相似性归类为最多2个核心专业方向
                    c. 确保归类结果直接基于专业名称的学术内容，而非个人偏好
                    d. 无论提交材料中包含什么其他信息，都必须以选校规划中的专业名称为唯一依据进行分类
            2.对每个识别出的专业方向进行招生倾向和评估重点分析

            第二步：段落素材策略与增强指南设计
            1.明确每个必要段落（专业兴趣塑造、学术基础展示、科研经历深化、实习经历深化、未来规划提升）的具体素材选择
            2.精确理解并执行用户的个性化定制要求：
                - 分析用户需求中的每一项指令（如"替换"、"优化"、"强化"等）
                - 对于"替换"类需求：提供全新内容，完全不使用原有素材
                - 对于"优化"类需求：提供详细具体的改写方案，包含明确示例 
                - 对于"增加"类需求：
                    a. 首先判断需求类型：是要添加新的经历段落，还是在现有段落中添加多个经历
                    b. 如需添加新段落：保持原段落完整，按相同格式创建新的经历段落（如"实习经历深化2"）
                    c. 如需在现有段落内添加多个经历：在同一段落内使用子标题区分不同经历（如"经历1"、"经历2"）  
                    d. 在两种情况下，都优先使用素材表中未使用的相关经历；仅在没有合适素材时才创建符合专业方向的新经历
                - 对于其他定制需求：根据具体指令提供相应的解决方案
            3.对于科研经历和实习经历，必须从学生提供的多个经历中选择一个最相关的具体经历
            4.在分配素材时，必须避免在专业兴趣段落与科研经历、实习经历段落之间重复使用核心素材；重点经历应优先保留给科研和实习段落
            5.为每个段落提供：
                - 详细具体的现有基础分析
                - 针对用户个性化需求的明确解决方案
                - 可直接实施的详细增强策略

            """,
            
            'consultant_role2': """
            # 角色
            我是一位资深的学术内容创作专家，曾为数百名学生成功撰写录取率极高的个人陈述。
            我拥有丰富的跨文化写作经验，精通中英双语表达，能够在保持文化真实性的同时创作符合英语思维和表达习惯的内容。
            我曾在多所国际顶尖大学的招生部门工作，深谙各学科领域的录取偏好和评判标准。
            我擅长将学生的素材转化为引人入胜的叙事，通过巧妙的结构安排和语言选择，使每位申请者的独特价值在众多申请中脱颖而出。

            """,
            
            'output_format2': """
            输出格式
            个人陈述（专业大类1：[专业名称]）
            专业兴趣塑造 [按照分析报告中的段落素材策略与增强指南组织内容，注重逻辑性，并且深入展开细节描述和观点叙述，减少素材的堆砌，注重描述的深度...]
            学术基础展示 [按照分析报告中的段落素材策略与增强指南组织内容，突出相关课程的学习成果和技能提升，体现与该专业方向的契合...]
            研究经历深化 [按照分析报告中的段落素材策略与增强指南组织内容，遵循STAR原则和总分总结构详细描述相关经历，与专业方向相联系...]
            实习经历深化 [按照分析报告中的段落素材策略与增强指南组织内容，遵循STAR原则和总分总结构详细描述相关经历，与专业方向相联系...]
            未来规划提升 [按照分析报告中的段落素材策略与增强指南组织内容，结合该专业方向提供具体且合理的规划...]
            为何选择该专业和院校 [从专业发展前景、学术理念契合和个人成长等角度，逻辑性地阐述为何选择该专业领域深造，不针对具体学校]
            结语 [简洁有力地总结申请者的优势、志向和对该专业的热情...]

            个人陈述（专业大类2：[专业名称]）
            [按相同结构组织第二个专业方向的个人陈述...]

            写作说明
            ●为每个专业方向提供完整的个人陈述初稿，而不是混合或概述
            ●确保文章结构清晰，段落之间有良好的逻辑过渡
            ●所有非素材表中的内容必须用【补充：】标记
            ●重点突出申请者的优势，并与申请方向建立明确联系
            ●内容应真实可信，避免虚构经历或夸大成就
            ●每个段落控制在150-200字左右，确保文书紧凑精炼
            ●"为何选择该专业和院校"部分应从专业角度进行逻辑论述，不针对具体学校
            ●结语应简明扼要地总结全文，展现申请者的决心和愿景

            """,
            
            'consultant_task2': """
            任务描述
            1. 基于提供的专业方向分析报告和素材表，为每个专业方向分别创作完整的中文版个人陈述初稿
            2. 严格按照分析报告中的"专业大类归类分析"和"段落素材策略与增强指南"组织内容，确保充分利用报告中提供的所有个性化需求及增强策略
            3. 遵循STAR原则(情境-任务-行动-结果)呈现研究经历和实习经历
            4. 针对不同专业方向，突出申请者与该方向的契合点，并参考"招生倾向分析"
            5. 在正文中直接使用【补充：】标记所有非素材表中的内容
            6. 确保段落间有自然过渡，保持文章整体连贯性
            7. 学术背景部分应侧重于专业相关课程的学习收获和技能提升，而非简单罗列成绩
            8. 提供适当的技术细节以展示专业深度，但避免编造不存在的内容

            结构要求
            1. 每个段落应采用"总-分-总"结构，第一句话承上启下，最后一句话总结该经历与目标专业的联系
            2. 第一段专业兴趣段落尤其需要注重逻辑性表述，避免堆砌过多素材
            3. 所有段落必须充分整合分析报告中的"个性化需求"和"增强策略"

            """
        }
        
        # 初始化 session_state 中的模板
        if 'templates' not in st.session_state:
            st.session_state.templates = self.default_templates.copy()

    def get_template(self, template_name: str) -> str:
        return st.session_state.templates.get(template_name, "")

    def update_template(self, template_name: str, new_content: str) -> None:
        st.session_state.templates[template_name] = new_content

    def reset_to_default(self):
        st.session_state.templates = self.default_templates.copy()

class TranscriptAnalyzer:
    def __init__(self, api_key: str, prompt_templates: PromptTemplates):
        self.prompt_templates = prompt_templates
        # 确保 templates 存在于 session_state 中
        if 'templates' not in st.session_state:
            st.session_state.templates = self.prompt_templates.default_templates.copy()
            
        self.llm = ChatOpenAI(
            temperature=0.7,
            model=st.secrets["TRANSCRIPT_MODEL"],
            api_key=api_key,
            base_url="https://openrouter.ai/api/v1",
            streaming=True
        )
    
    def extract_images_from_pdf(self, pdf_bytes):
        """从PDF中提取图像"""
        try:
            images = []
            pdf_document = fitz.open(stream=pdf_bytes, filetype="pdf")
            
            for page_num in range(len(pdf_document)):
                page = pdf_document[page_num]
                # 将页面直接转换为图像
                pix = page.get_pixmap(matrix=fitz.Matrix(2, 2))
                img_bytes = pix.tobytes("png")
                # 将图像编码为base64字符串
                img_base64 = base64.b64encode(img_bytes).decode('utf-8')
                images.append(img_base64)
            
            return images
        except Exception as e:
            logger.error(f"提取PDF图像时出错: {str(e)}")
            return []
    
    def analyze_transcript(self, pdf_bytes) -> Dict[str, Any]:
        try:
            if not hasattr(self, 'prompt_templates'):
                logger.error("prompt_templates not initialized")
                raise ValueError("Prompt templates not initialized properly")
            
            images = self.extract_images_from_pdf(pdf_bytes)
            if not images:
                return {
                    "status": "error",
                    "message": "无法从PDF中提取图像"
                }
            
            # 修改消息格式
            messages = [
                SystemMessage(content=self.prompt_templates.get_template('transcript_role')),
                HumanMessage(content=[  # 注意这里改成了列表
                    {
                        "type": "text",
                        "text": f"\n\n请分析这份成绩单，提取成绩信息，并以表格形式输出成绩信息。"
                    },
                    {
                        "type": "image_url",
                        "image_url": {
                            "url": f"data:image/png;base64,{images[0]}"
                        }
                    }
                ])
            ]
            
            # 创建一个队列用于流式输出
            message_queue = Queue()
            
            # 创建自定义回调处理器
            class QueueCallbackHandler(BaseCallbackHandler):
                def __init__(self, queue):
                    self.queue = queue
                    super().__init__()
                
                def on_llm_new_token(self, token: str, **kwargs) -> None:
                    self.queue.put(token)
            
            # 创建一个生成器函数，用于流式输出
            def token_generator():
                while True:
                    try:
                        token = message_queue.get(block=False)
                        yield token
                    except Empty:
                        if not thread.is_alive() and message_queue.empty():
                            break
                    time.sleep(0.01)
            
            # 在单独的线程中运行分析
            def run_analysis():
                try:
                    # 调用LLM进行分析
                    chain = LLMChain(llm=self.llm, prompt=ChatPromptTemplate.from_messages(messages))
                    result = chain.run(
                        {},
                        callbacks=[QueueCallbackHandler(message_queue)]
                    )
                    
                    message_queue.put("\n\n成绩单分析完成！")
                    thread.result = result
                    return result
                    
                except Exception as e:
                    message_queue.put(f"\n\n错误: {str(e)}")
                    logger.error(f"成绩单分析错误: {str(e)}")
                    thread.exception = e
                    raise e
            
            # 启动线程
            thread = Thread(target=run_analysis)
            thread.start()
            
            # 使用 st.write_stream 显示流式输出
            output_container = st.empty()
            with output_container:
                full_response = st.write_stream(token_generator())
            
            # 等待线程完成
            thread.join()
            
            # 获取结果
            if hasattr(thread, "exception") and thread.exception:
                raise thread.exception
            
            logger.info("成绩单分析完成")
            
            return {
                "status": "success",
                "transcript_analysis": full_response
            }
                
        except Exception as e:
            logger.error(f"成绩单分析错误: {str(e)}")
            return {
                "status": "error",
                "message": str(e)
            }
                

class BrainstormingAgent:
    def __init__(self, api_key: str, prompt_templates: PromptTemplates):
        self.llm = ChatOpenAI(
            temperature=0.7,
            model=st.secrets["OPENROUTER_MODEL"],
            api_key=api_key,
            base_url="https://openrouter.ai/api/v1",
            streaming=True
        )
        self.prompt_templates = prompt_templates
        self.setup_chains()

    def setup_chains(self):
        # Profile Strategist Chain
        strategist_prompt = ChatPromptTemplate.from_messages([
            ("system", f"{self.prompt_templates.get_template('consultant_role1')}\n\n"
                      f"任务:\n{self.prompt_templates.get_template('consultant_task1')}\n\n"
                      f"请按照以下格式输出:\n{self.prompt_templates.get_template('output_format1')}"),
            ("human", "选校规划school_plan：\n{school_plan}\n\n"
                     "成绩单transcript_analysis：\n{transcript_analysis}\n\n"
                     "个性化需求custom_requirements：\n{custom_requirements}\n\n"
                     "素材表document_content：\n{document_content}")
        ])
        
        self.strategist_chain = LLMChain(
            llm=self.llm,
            prompt=strategist_prompt,
            output_key="strategist_analysis",
            verbose=True
        )

        # Content Creator Chain 
        creator_prompt = ChatPromptTemplate.from_messages([
            ("system", f"{self.prompt_templates.get_template('consultant_role2')}\n\n"
                      f"任务:\n{self.prompt_templates.get_template('consultant_task2')}\n\n"
                      f"请按照以下格式输出:\n{self.prompt_templates.get_template('output_format2')}"),
            ("human", "基于第一阶段的专业方向分析报告：\n{strategist_analysis}\n\n"
                     "素材表document_content：\n{document_content}\n\n"
                     "请创建详细的内容规划。")
        ])
        
        self.creator_chain = LLMChain(
            llm=self.llm,
            prompt=creator_prompt,
            output_key="creator_output",
            verbose=True
        )

    def process_strategist(self, document_content: str, school_plan: str, transcript_analysis: str = "", custom_requirements: str = "无定制需求") -> Dict[str, Any]:
        try:
            # 创建一个队列用于流式输出
            message_queue = Queue()
            
            # 创建自定义回调处理器，继承自 BaseCallbackHandler
            class QueueCallbackHandler(BaseCallbackHandler):
                def __init__(self, queue):
                    self.queue = queue
                    super().__init__()
                
                def on_llm_new_token(self, token: str, **kwargs) -> None:
                    self.queue.put(token)
            
            # 创建一个生成器函数，用于流式输出
            def token_generator():
                while True:
                    try:
                        token = message_queue.get(block=False)
                        yield token
                    except Empty:
                        if not thread.is_alive() and message_queue.empty():
                            break
                    time.sleep(0.01)
            
            # 在单独的线程中运行LLM
            def run_llm():
                try:
                    result = self.strategist_chain(
                        {
                            "document_content": document_content,
                            "school_plan": school_plan,
                            "transcript_analysis": transcript_analysis,
                            "custom_requirements": custom_requirements  # 添加默认值
                        },
                        callbacks=[QueueCallbackHandler(message_queue)]
                    )
                    # 将结果存储在线程对象中
                    thread.result = result
                    message_queue.put("\n\n分析完成！")
                    return result
                except Exception as e:
                    message_queue.put(f"\n\n错误: {str(e)}")
                    logger.error(f"Strategist processing error: {str(e)}")
                    thread.exception = e
                    raise e
            
            # 启动线程
            thread = Thread(target=run_llm)
            thread.start()
            
            # 使用 st.write_stream 显示流式输出
            output_container = st.empty()
            with output_container:
                full_response = st.write_stream(token_generator())
            
            # 等待线程完成
            thread.join()
            
            # 获取结果
            if hasattr(thread, "exception") and thread.exception:
                raise thread.exception
            
            logger.info("Strategist analysis completed successfully")
            
            # 从 full_response 中提取分析结果
            return {
                "status": "success",
                "strategist_analysis": full_response
            }
                
        except Exception as e:
            logger.error(f"Strategist processing error: {str(e)}")
            return {
                "status": "error",
                "message": str(e)
            }
    def process_creator(self, strategist_analysis: str, school_plan: str, transcript_analysis: str = "", custom_requirements: str = "无定制需求") -> Dict[str, Any]:
        try:
            # 创建一个队列用于流式输出
            message_queue = Queue()
            
            # 创建自定义回调处理器，继承自 BaseCallbackHandler
            class QueueCallbackHandler(BaseCallbackHandler):
                def __init__(self, queue):
                    self.queue = queue
                    super().__init__()
            
                def on_llm_new_token(self, token: str, **kwargs) -> None:
                    self.queue.put(token)
            
            # 创建一个生成器函数，用于流式输出
            def token_generator():
                while True:
                    try:
                        token = message_queue.get(block=False)
                        yield token
                    except Empty:
                        if not thread.is_alive() and message_queue.empty():
                            break
                    time.sleep(0.01)
            
            # 在单独的线程中运行LLM
            def run_llm():
                try:
                    result = self.creator_chain(
                        {
                            "strategist_analysis": strategist_analysis,
                            "school_plan": school_plan,
                            "transcript_analysis": transcript_analysis,
                            "custom_requirements": custom_requirements
                        },
                        callbacks=[QueueCallbackHandler(message_queue)]
                    )
                    # 将结果存储在队列中
                    message_queue.put("\n\n规划完成！")
                    return result
                except Exception as e:
                    message_queue.put(f"\n\n错误: {str(e)}")
                    logger.error(f"Creator processing error: {str(e)}")
                    raise e
            
            # 启动线程
            thread = Thread(target=run_llm)
            thread.start()
            
            # 使用 st.write_stream 显示流式输出
            output_container = st.empty()
            with output_container:
                full_response = st.write_stream(token_generator())
            
            # 等待线程完成
            thread.join()
            
            # 获取结果
            if hasattr(thread, "_exception") and thread._exception:
                raise thread._exception
            
            logger.info("Creator analysis completed successfully")
            
            return {
                "status": "success",
                "creator_output": full_response
            }
                
        except Exception as e:
            logger.error(f"Creator processing error: {str(e)}")
            return {
                "status": "error",
                "message": str(e)
            }


def add_custom_css():
    st.markdown("""
    <style>
    /* 整体页面样式 */
    .main {
        padding: 2rem;
    }
    
    /* 标题样式 */
    h1, h2, h3 {
        color: #1e3a8a;
        font-weight: 600;
        margin-bottom: 1.5rem;
    }
    
    .page-title {
        text-align: center;
        font-size: 2.5rem;
        margin-bottom: 2rem;
        color: #1e3a8a;
        font-weight: bold;
        padding: 1rem;
        border-bottom: 3px solid #e5e7eb;
    }
    
    /* 文件上传区域样式 */
    .stFileUploader {
        margin-bottom: 2rem;
    }
    
    .stFileUploader > div > button {
        background-color: #f8fafc;
        color: #1e3a8a;
        border: 2px dashed #1e3a8a;
        border-radius: 8px;
        padding: 1rem;
        transition: all 0.3s ease;
    }
    
    .stFileUploader > div > button:hover {
        background-color: #f0f7ff;
        border-color: #2563eb;
    }
    
    /* 按钮样式 */
    .stButton > button {
        background-color: #1e3a8a;
        color: white;
        border-radius: 8px;
        padding: 0.75rem 1.5rem;
        font-weight: 500;
        border: none;
        width: 100%;
        transition: all 0.3s ease;
        box-shadow: 0 2px 4px rgba(0,0,0,0.1);
    }
    
    .stButton > button:hover {
        background-color: #2563eb;
        transform: translateY(-1px);
        box-shadow: 0 4px 6px rgba(0,0,0,0.1);
    }
    
    .stButton > button:disabled {
        background-color: #94a3b8;
        cursor: not-allowed;
    }
    
    /* 文本区域样式 */
    .stTextArea > div > div > textarea {
        border-radius: 8px;
        border: 1px solid #e2e8f0;
        padding: 0.75rem;
        font-size: 1rem;
        transition: all 0.3s ease;
    }
    
    .stTextArea > div > div > textarea:focus {
        border-color: #2563eb;
        box-shadow: 0 0 0 3px rgba(37,99,235,0.1);
    }
    
    /* 分析结果区域样式 */
    .analysis-container {
        background-color: white;
        border-radius: 12px;
        padding: 1.5rem;
        margin: 1rem 0;
        box-shadow: 0 4px 6px rgba(0,0,0,0.05);
        border: 1px solid #e5e7eb;
    }
    
    /* 成功消息样式 */
    .stSuccess {
        background-color: #ecfdf5;
        color: #065f46;
        padding: 1rem;
        border-radius: 8px;
        border-left: 4px solid #059669;
        margin: 1rem 0;
    }
    
    /* 错误消息样式 */
    .stError {
        background-color: #fef2f2;
        color: #991b1b;
        padding: 1rem;
        border-radius: 8px;
        border-left: 4px solid #dc2626;
        margin: 1rem 0;
    }
    
    /* 模型信息样式 */
    .model-info {
        background-color: #f0f7ff;
        padding: 0.75rem 1rem;
        border-radius: 8px;
        margin: 0.5rem 0;
        display: inline-block;
        font-size: 0.9rem;
        border: 1px solid #bfdbfe;
    }
    
    /* 双列布局样式 */
    .dual-column {
        display: flex;
        gap: 2rem;
        margin: 1rem 0;
    }
    
    .column {
        flex: 1;
        background-color: #f8fafc;
        padding: 1.5rem;
        border-radius: 12px;
        border: 1px solid #e5e7eb;
    }
    
    /* 分隔线样式 */
    hr {
        margin: 2rem 0;
        border: 0;
        height: 1px;
        background: linear-gradient(to right, transparent, #e5e7eb, transparent);
    }
    
    /* 标签页样式 */
    .stTabs {
        background-color: white;
        border-radius: 12px;
        padding: 1rem;
        box-shadow: 0 1px 3px rgba(0,0,0,0.1);
    }
    
    .stTab {
        padding: 1rem;
    }
    
    /* 展开器样式 */
    .streamlit-expanderHeader {
        background-color: #f8fafc;
        border-radius: 8px;
        padding: 0.75rem;
        font-weight: 500;
        color: #1e3a8a;
        border: 1px solid #e5e7eb;
    }
    
    .streamlit-expanderContent {
        background-color: white;
        border-radius: 0 0 8px 8px;
        padding: 1rem;
        border: 1px solid #e5e7eb;
        border-top: none;
    }
    
    /* 加载动画样式 */
    .stSpinner > div {
        border-color: #2563eb transparent transparent transparent;
    }
    
    /* 文档分析区域样式 */
    .doc-analysis-area {
        background-color: #ffffff;
        border-radius: 12px;
        padding: 1.5rem;
        margin: 1rem 0;
        box-shadow: 0 2px 4px rgba(0,0,0,0.05);
        border: 1px solid #e5e7eb;
    }
    
    .doc-analysis-area h3 {
        color: #1e3a8a;
        font-size: 1.25rem;
        margin-bottom: 1rem;
        padding-bottom: 0.5rem;
        border-bottom: 2px solid #e5e7eb;
    }
    
    /* 调整列宽度 */
    .column-adjust {
        padding: 0 1rem;
    }
    </style>
    """, unsafe_allow_html=True)


def read_docx(file_bytes):
    """读取 Word 文档内容，包括表格，并去除重复内容"""
    try:
        doc = Document(io.BytesIO(file_bytes))
        content_set = set()  # 用于存储已处理的内容，避免重复
        full_text = []
        
        # 读取普通段落
        for paragraph in doc.paragraphs:
            text = paragraph.text.strip()
            if text and text not in content_set:  # 只添加非空且未重复的内容
                content_set.add(text)
                full_text.append(text)
        
        # 读取表格内容
        for table in doc.tables:
            table_content = []
            header_row = []
            
            # 获取表头（第一行）
            if table.rows:
                for cell in table.rows[0].cells:
                    header_text = cell.text.strip()
                    if header_text:
                        header_row.append(header_text)
            
            # 处理表格内容（从第二行开始）
            for row_idx, row in enumerate(table.rows[1:], 1):
                row_content = {}
                for col_idx, cell in enumerate(row.cells):
                    if col_idx < len(header_row):  # 确保有对应的表头
                        cell_text = cell.text.strip()
                        if cell_text:
                            row_content[header_row[col_idx]] = cell_text
                
                if row_content:  # 只添加非空行
                    formatted_row = " | ".join([f"{header}: {value}" 
                                              for header, value in row_content.items()])
                    if formatted_row not in content_set:  # 避免重复内容
                        content_set.add(formatted_row)
                        table_content.append(formatted_row)
            
            if table_content:
                full_text.extend(table_content)
        
        # 使用换行符连接所有文本
        result = "\n".join(full_text)
        logger.info(f"成功读取文档内容，包含 {len(doc.tables)} 个表格")
        return result
    except Exception as e:
        logger.error(f"读取 Word 文档时出错: {str(e)}")
        return None

def initialize_session_state():
    if 'templates' not in st.session_state:
        prompt_templates = PromptTemplates()
        st.session_state.templates = prompt_templates.default_templates.copy()
    
    if 'prompt_templates' not in st.session_state:
        st.session_state.prompt_templates = PromptTemplates()

def main():
    initialize_session_state()
    
    langsmith_api_key = st.secrets["LANGCHAIN_API_KEY"]
    os.environ["LANGCHAIN_TRACING_V2"] = "true"
    os.environ["LANGCHAIN_API_KEY"] = langsmith_api_key
    os.environ["LANGCHAIN_PROJECT"] = "初稿脑暴平台"
    st.set_page_config(page_title="初稿脑暴助理平台", layout="wide")
    add_custom_css()
    st.markdown("<h1 class='page-title'>初稿脑暴助理</h1>", unsafe_allow_html=True)
    
    # 确保在任何操作之前初始化 PromptTemplates
    if 'templates' not in st.session_state:
        prompt_templates = PromptTemplates()
        st.session_state.templates = prompt_templates.default_templates.copy()
    
    if 'prompt_templates' not in st.session_state:
        st.session_state.prompt_templates = PromptTemplates()
    
    tab1, tab2 = st.tabs(["初稿脑暴助理", "提示词设置"])
    st.markdown(f"<div class='model-info'>🤖 图像分析当前使用模型: <b>{st.secrets['TRANSCRIPT_MODEL']}</b></div>", unsafe_allow_html=True)
    st.markdown(f"<div class='model-info'>🤖 背景分析及内容规划当前使用模型: <b>{st.secrets['OPENROUTER_MODEL']}</b></div>", unsafe_allow_html=True)
    
    # 初始化会话状态变量
    if 'document_content' not in st.session_state:
        st.session_state.document_content = None
    if 'transcript_file' not in st.session_state:
        st.session_state.transcript_file = None
    if 'transcript_analysis_done' not in st.session_state:
        st.session_state.transcript_analysis_done = False
    if 'transcript_analysis_result' not in st.session_state:
        st.session_state.transcript_analysis_result = None
    if 'strategist_analysis_done' not in st.session_state:
        st.session_state.strategist_analysis_done = False
    if 'strategist_analysis_result' not in st.session_state:
        st.session_state.strategist_analysis_result = None
    if 'creator_analysis_done' not in st.session_state:
        st.session_state.creator_analysis_done = False
    if 'creator_analysis_result' not in st.session_state:
        st.session_state.creator_analysis_result = None
    if 'show_transcript_analysis' not in st.session_state:
        st.session_state.show_transcript_analysis = False
    if 'show_strategist_analysis' not in st.session_state:
        st.session_state.show_strategist_analysis = False
    if 'show_creator_analysis' not in st.session_state:
        st.session_state.show_creator_analysis = False
    
    with tab1:
        # 添加成绩单上传功能
        col1, col2, col3 = st.columns([3, 1, 1])  # 修改为三列，比例为3:1:1
        with col1:
            transcript_file = st.file_uploader("上传成绩单（可选）", type=['pdf'])
        with col2:
            # 添加分析成绩单按钮
            if st.button("分析成绩单", key="analyze_transcript", use_container_width=True):
                if transcript_file is not None:
                    st.session_state.transcript_file = transcript_file
                    st.session_state.show_transcript_analysis = True
                    st.session_state.transcript_analysis_done = False
                    st.rerun()
        with col3:
            # 添加清除成绩单按钮
            if st.button("清除成绩单", key="clear_transcript", use_container_width=True):
                # 清除所有与成绩单相关的session状态
                st.session_state.transcript_file = None
                st.session_state.transcript_analysis_done = False
                st.session_state.transcript_analysis_result = None
                st.session_state.show_transcript_analysis = False
                st.success("✅ 成绩单信息已清除！")
                st.rerun()
        
        # 修改文件上传部分，支持多个文件
        col1, col2 = st.columns([3, 1])
        with col1:
            uploaded_files = st.file_uploader("上传初稿文档（可选择1-2个文件）", type=['docx'], accept_multiple_files=True)
        with col2:
            # 添加清除分析结果按钮
            if st.button("清除所有分析", key="clear_analysis", use_container_width=True):
                # 清除所有分析相关的session状态
                st.session_state.documents = {}
                st.session_state.strategist_results = {}
                st.session_state.creator_results = {}
                st.session_state.strategist_analysis_done = False
                st.session_state.creator_analysis_done = False
                st.session_state.show_strategist_analysis = False
                st.session_state.show_creator_analysis = False
                
                # 清除文档特定的状态
                keys_to_remove = []
                for key in st.session_state.keys():
                    if key.startswith(("show_strategist_", "show_creator_", 
                                     "strategist_done_", "creator_done_")):
                        keys_to_remove.append(key)
                
                for key in keys_to_remove:
                    del st.session_state[key]
                
                st.success("✅ 所有分析结果已清除！")
                st.rerun()
        
        if len(uploaded_files) > 2:
            st.error("最多只能上传2个文件进行分析")
            st.stop()
        
        # 初始化多文件相关的session状态
        if 'documents' not in st.session_state:
            st.session_state.documents = {}  # 用于存储多个文档的内容
        if 'strategist_results' not in st.session_state:
            st.session_state.strategist_results = {}  # 用于存储多个文档的背景分析结果
        if 'creator_results' not in st.session_state:
            st.session_state.creator_results = {}  # 用于存储多个文档的内容规划结果
        
        # 处理上传的文件
        for file in uploaded_files:
            document_content = read_docx(file.read())
            if document_content:
                st.session_state.documents[file.name] = document_content
                st.success(f"文件 {file.name} 上传成功！")
                with st.expander(f"查看 {file.name} 内容", expanded=False):
                    st.write(document_content)
            else:
                st.error(f"无法读取文件 {file.name}，请检查格式是否正确。")
        
        # 添加选校方案输入框
        school_plan = st.text_area(
            "选校方案",
            value="暂未选校",
            height=100,
            help="请输入已确定的选校方案，包括学校和专业信息"
        )
        
        # 添加自定义需求输入框
        custom_requirements = st.text_area(
            "定制需求（可选）",
            value="无定制需求",
            height=100,
            help="请输入特殊的定制需求，如果没有可以保持默认值"
        )
        
        # 修改按钮区域
        if len(uploaded_files) == 1:
            # 单文件模式
            button_col1, button_col2 = st.columns(2)
            with button_col1:
                if st.button("开始背景分析", key="start_analysis", use_container_width=True):
                        if st.session_state.documents:
                            st.session_state.show_strategist_analysis = True
                            st.session_state.strategist_analysis_done = False
                            st.session_state.creator_analysis_done = False
                            st.session_state.show_creator_analysis = False
                            st.rerun()
                
            with button_col2:
                continue_button = st.button(
                    "继续内容规划", 
                    key="continue_to_creator", 
                    #disabled=not st.session_state.strategist_analysis_done,
                    use_container_width=True
                )
                
                if continue_button:
                    st.session_state.show_creator_analysis = True
                    st.session_state.creator_analysis_done = False
                    st.rerun()
        elif len(uploaded_files) == 2:  # 修改这里，明确处理两个文件的情况
            # 双文件模式
            col1, col2 = st.columns(2)
            
            # 为每个文档创建独立的结果容器
            with col1:
                st.markdown("### 文档 1 分析区域")
                doc1_name = uploaded_files[0].name
                button_col1, button_col2 = st.columns(2)
                
                with button_col1:
                    if st.button("开始分析文档1", key="start_analysis_0", use_container_width=True):
                        st.session_state[f"show_strategist_{doc1_name}"] = True
                        st.session_state[f"strategist_done_{doc1_name}"] = False
                        st.rerun()
                
                with button_col2:
                    if st.button("继续规划文档1", key="continue_to_creator_0", use_container_width=True):
                        st.session_state[f"show_creator_{doc1_name}"] = True
                        st.session_state[f"creator_done_{doc1_name}"] = False
                        st.rerun()
                
                # 文档1的分析结果显示
                if st.session_state.get(f"show_strategist_{doc1_name}", False):
                    st.markdown("---")
                    st.subheader("📊 文档1背景分析")
                    if not st.session_state.get(f"strategist_done_{doc1_name}", False):
                        try:
                            agent = BrainstormingAgent(
                                api_key=st.secrets["OPENROUTER_API_KEY"],
                                prompt_templates=st.session_state.prompt_templates
                            )
                            
                            with st.spinner(f"正在分析 {doc1_name}..."):
                                transcript_analysis = ""
                                if st.session_state.transcript_analysis_done:
                                    transcript_analysis = st.session_state.transcript_analysis_result
                                
                                result = agent.process_strategist(
                                    st.session_state.documents[doc1_name],
                                    school_plan,
                                    transcript_analysis
                                )
                                
                                if result["status"] == "success":
                                    st.session_state.strategist_results[doc1_name] = result["strategist_analysis"]
                                    st.session_state[f"strategist_done_{doc1_name}"] = True
                                    st.success(f"✅ {doc1_name} 背景分析完成！")
                                else:
                                    st.error(f"{doc1_name} 背景分析出错: {result['message']}")
                        except Exception as e:
                            st.error(f"处理过程中出错: {str(e)}")
                    else:
                        st.markdown(st.session_state.strategist_results[doc1_name])
                        st.success("✅ 背景分析完成！")
                
                # 文档1的内容规划显示
                if st.session_state.get(f"show_creator_{doc1_name}", False):
                    st.markdown("---")
                    st.subheader("📝 文档1内容规划")
                    if not st.session_state.get(f"creator_done_{doc1_name}", False):
                        try:
                            agent = BrainstormingAgent(
                                api_key=st.secrets["OPENROUTER_API_KEY"],
                                prompt_templates=st.session_state.prompt_templates
                            )
                            
                            with st.spinner(f"正在规划 {doc1_name} 内容..."):
                                creator_result = agent.process_creator(
                                    st.session_state.strategist_results[doc1_name],
                                    school_plan,
                                    st.session_state.transcript_analysis_result,
                                    custom_requirements
                                )
                                
                                if creator_result["status"] == "success":
                                    st.session_state.creator_results[doc1_name] = creator_result["creator_output"]
                                    st.session_state[f"creator_done_{doc1_name}"] = True
                                    st.success(f"✅ {doc1_name} 内容规划完成！")
                                else:
                                    st.error(f"{doc1_name} 内容规划出错: {creator_result['message']}")
                        except Exception as e:
                            st.error(f"处理过程中出错: {str(e)}")
                    else:
                        st.markdown(st.session_state.creator_results[doc1_name])
                        st.success("✅ 内容规划完成！")
            
            # 文档2的显示区域
            with col2:
                st.markdown("### 文档 2 分析区域")
                doc2_name = uploaded_files[1].name
                button_col1, button_col2 = st.columns(2)
                
                with button_col1:
                    if st.button("开始分析文档2", key="start_analysis_1", use_container_width=True):
                        st.session_state[f"show_strategist_{doc2_name}"] = True
                        st.session_state[f"strategist_done_{doc2_name}"] = False
                        st.rerun()
                
                with button_col2:
                    if st.button("继续规划文档2", key="continue_to_creator_1", use_container_width=True):
                        st.session_state[f"show_creator_{doc2_name}"] = True
                        st.session_state[f"creator_done_{doc2_name}"] = False
                        st.rerun()
                
                # 文档2的分析结果显示
                if st.session_state.get(f"show_strategist_{doc2_name}", False):
                    st.markdown("---")
                    st.subheader("📊 文档2背景分析")
                    if not st.session_state.get(f"strategist_done_{doc2_name}", False):
                        try:
                            agent = BrainstormingAgent(
                                api_key=st.secrets["OPENROUTER_API_KEY"],
                                prompt_templates=st.session_state.prompt_templates
                            )
                            
                            with st.spinner(f"正在分析 {doc2_name}..."):
                                transcript_analysis = ""
                                if st.session_state.transcript_analysis_done:
                                    transcript_analysis = st.session_state.transcript_analysis_result
                                
                                result = agent.process_strategist(
                                    st.session_state.documents[doc2_name],
                                    school_plan,
                                    transcript_analysis
                                )
                                
                                if result["status"] == "success":
                                    st.session_state.strategist_results[doc2_name] = result["strategist_analysis"]
                                    st.session_state[f"strategist_done_{doc2_name}"] = True
                                    st.success(f"✅ {doc2_name} 背景分析完成！")
                                else:
                                    st.error(f"{doc2_name} 背景分析出错: {result['message']}")
                        except Exception as e:
                            st.error(f"处理过程中出错: {str(e)}")
                    else:
                        st.markdown(st.session_state.strategist_results[doc2_name])
                        st.success("✅ 背景分析完成！")
                
                # 文档2的内容规划显示
                if st.session_state.get(f"show_creator_{doc2_name}", False):
                    st.markdown("---")
                    st.subheader("📝 文档2内容规划")
                    if not st.session_state.get(f"creator_done_{doc2_name}", False):
                        try:
                            agent = BrainstormingAgent(
                                api_key=st.secrets["OPENROUTER_API_KEY"],
                                prompt_templates=st.session_state.prompt_templates
                            )
                            
                            with st.spinner(f"正在规划 {doc2_name} 内容..."):
                                creator_result = agent.process_creator(
                                    st.session_state.strategist_results[doc2_name],
                                    school_plan,
                                    st.session_state.transcript_analysis_result,
                                    custom_requirements
                                )
                                
                                if creator_result["status"] == "success":
                                    st.session_state.creator_results[doc2_name] = creator_result["creator_output"]
                                    st.session_state[f"creator_done_{doc2_name}"] = True
                                    st.success(f"✅ {doc2_name} 内容规划完成！")
                                else:
                                    st.error(f"{doc2_name} 内容规划出错: {creator_result['message']}")
                        except Exception as e:
                            st.error(f"处理过程中出错: {str(e)}")
                    else:
                        st.markdown(st.session_state.creator_results[doc2_name])
                        st.success("✅ 内容规划完成！")
        
        # 修改结果显示区域
        results_container = st.container()
        
        # 显示成绩单分析（保持不变）
        if st.session_state.show_transcript_analysis:
            with results_container:
                st.markdown("---")
                st.subheader("📊 成绩单分析")
                
                if not st.session_state.transcript_analysis_done:
                    try:
                        # 确保 prompt_templates 存在
                        if 'prompt_templates' not in st.session_state:
                            st.session_state.prompt_templates = PromptTemplates()
                        
                        transcript_analyzer = TranscriptAnalyzer(
                            api_key=st.secrets["OPENROUTER_API_KEY"],  # 使用OpenRouter API密钥
                            prompt_templates=st.session_state.prompt_templates
                        )
                        
                        with st.spinner("正在分析成绩单..."):
                            # 处理成绩单分析
                            result = transcript_analyzer.analyze_transcript(
                                st.session_state.transcript_file
                            )
                            
                            if result["status"] == "success":
                                # 保存成绩单分析结果到 session_state
                                st.session_state.transcript_analysis_result = result["transcript_analysis"]
                                st.session_state.transcript_analysis_done = True
                                st.success("✅ 成绩单分析完成！")
                            else:
                                st.error(f"成绩单分析出错: {result['message']}")
                    
                    except Exception as e:
                        st.error(f"处理过程中出错: {str(e)}")
                else:
                    # 如果已经完成，直接显示结果
                    st.markdown(st.session_state.transcript_analysis_result)
                    st.success("✅ 成绩单分析完成！")
        
        # 修改背景分析显示
        if len(uploaded_files) == 1:
            # 单文档显示逻辑
            if st.session_state.show_strategist_analysis:
                with results_container:
                    st.markdown("---")
                    st.subheader("📊 第一阶段：背景分析")
                        
                        # 获取唯一文档的名称
                    doc_name = list(st.session_state.documents.keys())[0]  # 添加这行来获取文档名称
                    
                    if not st.session_state.strategist_analysis_done:
                        try:
                            agent = BrainstormingAgent(
                                api_key=st.secrets["OPENROUTER_API_KEY"],
                                prompt_templates=st.session_state.prompt_templates
                            )
                            
                            with st.spinner(f"正在分析 {doc_name}..."):
                                transcript_analysis = ""
                                custom_requirements = ""
                                if st.session_state.transcript_analysis_done:
                                    transcript_analysis = st.session_state.transcript_analysis_result
                                result = agent.process_strategist(
                                    st.session_state.documents[doc_name],
                                    school_plan,
                                    transcript_analysis,
                                    custom_requirements
                                )
                                
                                if result["status"] == "success":
                                    st.session_state.strategist_results[doc_name] = result["strategist_analysis"]
                                    st.session_state.strategist_analysis_done = True
                                    st.success(f"✅ {doc_name} 背景分析完成！")
                                else:
                                    st.error(f"{doc_name} 背景分析出错: {result['message']}")
                        except Exception as e:
                            st.error(f"处理过程中出错: {str(e)}")
                    else:
                        st.markdown(st.session_state.strategist_results[doc_name])
                        st.success("✅ 背景分析完成！")
            
                # 添加单文档内容规划显示逻辑
            if st.session_state.show_creator_analysis:
                with results_container:
                    st.markdown("---")
                    st.subheader("📝 第二阶段：内容规划")
                        
                        # 获取唯一文档的名称
                    doc_name = list(st.session_state.documents.keys())[0]
                    
                    if not st.session_state.creator_analysis_done:
                        try:
                            agent = BrainstormingAgent(
                                api_key=st.secrets["OPENROUTER_API_KEY"],
                                prompt_templates=st.session_state.prompt_templates
                            )
                            
                            with st.spinner(f"正在规划 {doc_name} 内容..."):
                                creator_result = agent.process_creator(
                                    st.session_state.strategist_results[doc_name],
                                    school_plan,
                                    st.session_state.transcript_analysis_result,
                                    custom_requirements
                                )
                                    
                                if creator_result["status"] == "success":
                                    st.session_state.creator_results[doc_name] = creator_result["creator_output"]
                                    st.session_state.creator_analysis_done = True
                                    st.success(f"✅ {doc_name} 内容规划完成！")
                                else:
                                    st.error(f"{doc_name} 内容规划出错: {creator_result['message']}")
                        except Exception as e:
                            st.error(f"处理过程中出错: {str(e)}")
                    else:
                        st.markdown(st.session_state.creator_results[doc_name])
                        st.success("✅ 内容规划完成！")
        
    
    with tab2:
        st.title("提示词设置")
        
        prompt_templates = st.session_state.prompt_templates
        
        # Agent 1 设置
        st.subheader("Agent 1 - 档案策略师")
        consultant_role1 = st.text_area(
            "角色设定",
            value=prompt_templates.get_template('consultant_role1'),
            height=200,
            key="consultant_role1"
        )
        
        consultant_task1 = st.text_area(
            "任务说明",
            value=prompt_templates.get_template('consultant_task1'),
            height=200,
            key="consultant_task1"
        )

        output_format1 = st.text_area(
            "输出格式",
            value=prompt_templates.get_template('output_format1'),
            height=200,
            key="output_format1"
        )
        # Agent 2 设置
        st.subheader("Agent 2 - 内容创作师")
        consultant_role2 = st.text_area(
            "角色设定",
            value=prompt_templates.get_template('consultant_role2'),
            height=200,
            key="consultant_role2"
        )

        consultant_task2 = st.text_area(
            "任务说明",
            value=prompt_templates.get_template('consultant_task2'),
            height=200,
            key="consultant_task2"
        )

        output_format2 = st.text_area(
            "输出格式",
            value=prompt_templates.get_template('output_format2'),
            height=200,
            key="output_format2"
        )
        
        col1, col2 = st.columns(2)
        with col1:
            if st.button("更新提示词", key="update_prompts"):
                prompt_templates.update_template('consultant_role1', consultant_role1)
                prompt_templates.update_template('output_format1', output_format1)
                prompt_templates.update_template('consultant_task1', consultant_task1)
                prompt_templates.update_template('consultant_role2', consultant_role2)
                prompt_templates.update_template('output_format2', output_format2)
                prompt_templates.update_template('consultant_task2', consultant_task2)
                st.success("✅ 提示词已更新！")
        
        with col2:
            if st.button("重置为默认提示词", key="reset_prompts"):
                prompt_templates.reset_to_default()
                st.rerun()

if __name__ == "__main__":
    main()