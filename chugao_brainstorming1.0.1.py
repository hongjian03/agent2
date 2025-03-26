import streamlit as st
import json
from langchain_openai import ChatOpenAI
from langchain.prompts import ChatPromptTemplate
from langchain.schema import HumanMessage, SystemMessage
from langchain.callbacks.streaming_stdout import StreamingStdOutCallbackHandler
from langchain.callbacks.streamlit import StreamlitCallbackHandler
from langchain.chains import SequentialChain, LLMChain
import os
from typing import Dict, Any, List
import logging
import sys
from docx import Document
import io
import base64
from PyPDF2 import PdfReader
from PIL import Image
import fitz  # PyMuPDF
# 配置日志记录
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(levelname)s - %(message)s',
    handlers=[
        logging.StreamHandler(sys.stdout)
    ]
)
from queue import Queue
from threading import Thread
import time
from queue import Empty
logger = logging.getLogger(__name__)
from langchain.callbacks.base import BaseCallbackHandler

# 记录程序启动
logger.info("程序开始运行")

# 只在第一次运行时替换 sqlite3
if 'sqlite_setup_done' not in st.session_state:
    try:
        logger.info("尝试设置 SQLite")
        __import__('pysqlite3')
        sys.modules['sqlite3'] = sys.modules.pop('pysqlite3')
        st.session_state.sqlite_setup_done = True
        logger.info("SQLite 设置成功")
    except Exception as e:
        logger.error(f"SQLite 设置错误: {str(e)}")
        st.session_state.sqlite_setup_done = True


class PromptTemplates:
    def __init__(self):
        # 定义示例数据作为字符串
        self.default_templates = {
            'transcript_role': """
            # 角色
            你是专业的成绩单分析师，擅长从成绩单中提取关键信息并以表格形式展示成绩。
            """,
            
            'transcript_task': """

            """,
            
            'transcript_output': """

            """,
            
            'consultant_role1': """
            # 角色
            我是一位拥有十年经验的顶尖大学招生文书评估专家。
            我擅长分析申请专业关联性，帮助学生找到最有竞争力的申请文书策略，减少重复工作并提高申请效率。
            """,
            
            'output_format1': """
            输出格式
            必须完整分析并输出所有识别出的专业方向（最多2个），每个方向的分析必须包含以下内容：
            
            个人陈述初稿写作策略报告
            
            一、专业方向分析
            专业方向1：[名称]
            所属院校及专业清单：
            ● 招生倾向分析：
            ● [招生倾向点1]  
            ● [招生倾向点2]  
            ● [招生倾向点3]  
            ● [...]  
            
            专业方向2：[名称] (如识别出第二个专业方向，必须完整分析)
            所属院校及专业清单：
            ● 招生倾向分析：
            ● [招生倾向点1]  
            ● [招生倾向点2]  
            ● [招生倾向点3]  
            ● [...]
            
            二、段落素材策略与增强指南
            必须为每个专业方向分别提供完整的段落策略分析：
            
            专业方向1
            1. 专业兴趣段落
            ● 选择素材：[选择的具体素材]  

            ● 现有基础：[简述现有素材]  

            ● 个性化需求：[根据用户提供的个性化需求分析判断，如无则标注"无特殊要求"]  

            ● 增强策略： 
            - A. 确定写作角度（必须选择一个主要角度并深入展开）： 
            - [选择一个最合适的角度：过去经历/时事新闻/科研成果]  
            - [阐述选择该角度的原因]  
            - B. 详细展开核心信息点：
            - [描述关键经历/事件/问题]  
            - [分析该信息点与专业方向的直接联系]  
            - [展示个人独特的思考和见解]  
            - C. 展示专业理解：
            - [阐述对申请专业的深度理解]  
            - [分析该专业的价值和影响力]  
            - [说明为何被该领域吸引]  
            - D. 建立个人联系：
            - [过去经历如何引导到申请专业]  
            - [个人特质与专业需求的契合点]  
            - [未来希望在该专业探索的方向]  

            2. 学术基础展示
            ● 选择素材：[选择的具体素材]

            ● 现有基础：[简述现有素材]

            ● 个性化需求：[根据用户提供的个性化需求分析判断，如无则标注"无特殊要求"]

            ● 增强策略：
            - A. 综述部分：
            - [说明本科课程如何为研究生申请打下基础]
            - [强调课程与申请专业的相关性]
            - [体现学习态度与专业潜力]
            - B. 学术亮点1：
            - [与申请专业相关的知识框架联系]
            - [具体相关课程/项目名称]
            - [课程/项目亮点内容简述]
            - [与申请专业的连接]
            - C. 学术亮点2：
            - [提升的专业相关学术能力]
            - [具体相关课程/项目名称]
            - [课程/项目亮点内容简述]
            - [与申请专业的连接]
            - D. 学术亮点3：
            - [提升的专业相关学术能力]
            - [具体相关课程/项目名称]
            - [课程/项目内容简述]
            - [与申请专业的连接]
            - E. 总结部分：
            - [强调已做好学术准备]
            - [表达对未来学习的期待]
            - [与申请专业建立明确的联系] 

            3. 科研经历深化
            ● 选择素材：[选择的具体素材]

            ● 现有基础：[简述现有素材]

            ● 选择理由：[阐述为何选择此素材而非其他素材，说明此经历相比其他可选经历的优势]

            ● 个性化需求：[根据用户提供的个性化需求分析判断，如无则标注"无特殊要求"]
            ● 增强策略：
            - A.明确项目背景（仅基于素材表中的具体描述）：
            - [详细说明项目的具体情况，不添加未提及的细节]
            - [解释研究主题选择原因，基于素材表中的实际信息]
            - B.深化研究过程（禁止添加未在素材表中提及的研究方法或技术）：
            - [描述研究方法学，仅限于素材表中明确提及的内容]
            - [增加研究中遇到的挑战及解决方法，必须有素材支持]
            - [添加与导师/专家交流内容，仅限于素材表中实际提及的交流]
            - C.强化研究成果（禁止虚构或夸大研究成果）：
            - [明确研究发现和建议，仅基于素材表中的实际结果]
            - [将研究与专业方向关联，不夸大影响力]
            - [分析研究对专业理解的影响，基于实际经历]
            - D.联系申请专业（基于实际研究经历）：
            - [分析研究如何启发对该专业的兴趣，基于事实]
            - [强调研究方法与专业的相关性，不添加未使用的方法] 

            4. 实习经历深化
            ● 选择素材：[必须只选择一个最相关的实习经历]

            ● 现有基础：[简述现有素材]

            ● 选择理由：[阐述为何选择此素材而非其他素材，说明此经历相比其他可选经历的优势]

            ● 个性化需求：[根据用户提供的个性化需求分析判断，如无则标注"无特殊要求"]
            ● 增强策略：
            - A.明确项目结构（仅限于素材表中描述的实际项目）：
            - [详细说明项目的具体目标和框架，不扩大项目范围]
            - [描述项目如何与专业方向相关，基于实际情况]
            - B.强化个人贡献（严禁虚构或夸大职责和成就）：
            - [具体阐述职责，使用ACTION-RESULT模式，仅限于素材表中提及的职责]
            - [量化成就，必须有素材表中明确数据支持，禁止自行创造数据]
            - [详细描述参与的具体案例，仅限于素材表中提及的案例]
            - C.深化专业见解（基于实际工作经历）：
            - [分析工作中面临的专业挑战，必须基于素材表中提及的实际挑战]
            - [反思专业相关问题，不添加未经历的问题]
            - [展示对行业理解，仅基于实际接触的行业知识]
            - D.与申请专业联系（基于真实经历）：
            - [说明经历如何引导对专业的兴趣，基于事实]
            - [解释为何需要该专业的系统知识，不夸大需求]
            - [表明经历如何确认了专业需求，基于实际体验]

            5. 未来规划提升
            ● 选择素材：[选择的具体素材]

            ● 现有基础：[简述现有素材]

            ● 个性化需求：[根据用户提供的个性化需求分析判断，如无则标注"无特殊要求"]

            ● 增强策略：
            - A.明确短期目标：
            - [详细说明希望获得的具体技能和知识]
            - [计划获取的特定认证或参与的项目]
            - [希望建立的学术和专业网络]
            - B.设计中期职业路径：
            - [毕业后1-3年的具体职业目标]
            - [希望解决的专业相关挑战]
            - [工作地理区域意向]
            - C.制定长期愿景：
            - [5-10年职业发展愿景]
            - [希望在专业领域的具体贡献]
            - [与申请专业训练的关联]
            - D.增加可行性：
            - [如何利用现有资源和网络]
            - [毕业后的具体目标组织或项目]
            - [展示对行业趋势的了解]

            专业方向2 (如识别出第二个专业方向，必须按照相同格式完整分析)
            [与专业方向1相同的段落结构] 

            重要说明：
            1.必须严格按照上述格式输出分析报告，包括分行的要求
            2.如果用户提供明确的申请专业方向，则必须直接使用且不进行任何修改或补充，绝不从素材表中重新解读或与素材表中的信息结合；如提供选校规划，则需进行分析和归类
            3.当用户明确提供申请方向时，不得基于素材表中的任何信息替换或修改该方向
            4.如果识别出两个专业方向，必须为两个方向分别提供完整的分析报告，且包含所有要点
            5.科研经历、实习经历必须只选择一个最相关的经历，不得选择多个
            6.所有分析必须基于学生提供的真实材料
            7.所有"个性化需求"必须根据用户实际提供的需求进行分析判断，如无则标注"无特殊要求"
            8.对于"增加经历"类需求，必须首先检查是否有未使用的相关素材，只有在没有可用素材时才创建新内容
            9.增强策略必须按照固定格式分点展示，确保清晰易读
            10.根据用户的实际需求灵活调整格式，确保所有增加的经历都被有效整合到相应的段落中，每个经历都应有完整的分析结构
            11.禁止输出任何未在上述格式中明确指定的内容，不要添加额外的解释、说明或任何形式的补充内容
            12.所有增强策略必须遵循严格的事实性原则：
            - 只能基于素材表中的实际信息提供建议，不得过度编撰数据或夸大经历细节
            - 所有建议必须符合申请人的实际学术能力和专业水平，避免提出申请人无法支撑的内容建议
            - 对于科研和实习经历，必须100%基于素材表中的具体描述，禁止添加未在素材表中明确提及的研究方法、技术、成果或职责
            - 量化成果必须有素材表中的明确数据支持，禁止自行设定或创造数据
            - 违反上述规则将导致文书真实性严重受损，必须严格遵守
            13.完成输出后不得增加任何总结性语句或询问用户是否满意的内容
            14.不得在格式中添加或修改任何标题、副标题或其他未在上述格式中明确列出的内容
            15.以markdown形式输出，确保经过markdown渲染后可以正确清晰的显示

            输出时必须按以下格式使用markdown标记:
            - 主标题使用 # (如：# 个人陈述初稿写作策略报告)
            - 二级标题使用 ## (如：## 一、专业方向分析)
            - 三级标题使用 ### (如：### 专业方向1：[名称])
            - 四级标题使用 #### (如：#### 1. 专业兴趣段落)
            - 列表项使用 * 或 - 并确保正确缩进子列表
            - 所有小标题如"选择素材"、"现有基础"等使用**加粗**处理
            - 重点内容使用*斜体*标记
            - 段落之间保留空行确保正确渲染
            - 专业方向名称、院校名称加粗显示
            - 确保符号与文字之间有空格(如"##" 和标题文字之间)

            完整示例：
            # 个人陈述初稿写作策略报告

            ## 一、专业方向分析
            ### 专业方向1：**[名称]**
            所属院校及专业清单：
            * 招生倾向分析：
              * [招生倾向点1]  
              * [招生倾向点2]  

            ## 二、段落素材策略与增强指南
            ### 专业方向1
            #### 1. 专业兴趣段落
            * **选择素材**：[选择的具体素材]
            
            * **现有基础**：[简述现有素材]
            
            * **个性化需求**：[需求分析]
            
            * **增强策略**：
              * A. 确定写作角度：
                * [写作角度]
                * [选择理由]
            """,
            
            'consultant_task1': """
            基本任务
            1.仔细阅读和充分理解用户上传的素材表document_content、输入的个性化需求custom_requirements、分析用户上传的成绩单transcript_analysis、分析用户提供的申请专业方向或详细的选校规划school_plan
            2.准确识别用户指定的申请方向：
            - 如用户提供明确的申请方向名称（如"公共政策"），则必须【严格并唯一】以用户提供的名称为【唯一】分析对象，【完全忽略素材表中的任何专业信息和方向名称】，不得结合或修改用户指定的方向名称
            - 如用户提供详细选校规划，则需分析各院校专业的共性与差异，归类为1-2个专业大类方向
            - 禁止根据用户提供的申请专业方向及心仪院校确认申请方向
            - 禁止将用户提供的素材表中的任何申请院校信息内容纳入选校规划的分析
            3.为每个识别出的专业大类设计一套完整的基础文书策略，包含专业兴趣塑造、学术基础展示、科研经历深化、实习经历深化、未来规划提升五个必要段落
            4.评估现有素材对各专业大类的适用性，提出具体可执行的优化建议

            重要规则
            1.禁止提及任何具体学校名称，即使素材表中提及了相关学校和专业的内容
            2.学术基础展示部分必须详细分析用户上传的成绩单transcript_analysis内容，【必须完全避免】讨论任何成绩相关信息，包括但不限于具体课程分数、GPA、课程平均分数值、百分比成绩，【仅】重点关注与申请专业相关的课程表现和学术能力，不讨论具体成绩表现当要求
            3."增加经历"时，应遵循以下原则：
            - 如果用户需求是增加新的经历段落，则保持原有段落不变，并添加一个新的完整段落，使用"XX经历深化1"、"XX经历深化2"等标题格式
            - 如果用户需求是在现有段落中增加多个经历，则在同一段落内使用子标题"XX经历1"、"XX经历2"等进行区分
            - 增加的经历段落或子段落必须保持与其他段落相同的结构和深度
            4.此规则适用于所有类型的经历（科研、实习、课外活动等），每个经历段落或子段落都必须有完整的分析结构
            5.成绩单缺失情况下的处理原则：
            - 学术基础展示段落必须基于学生在读专业及相关素材进行分析
            - 重点关注素材表中申请专业与学生当前专业的衔接关系
            - 分析过往相关教育经历中体现的学术优势
            - 重点挖掘研究经历中的课堂作业部分，作为学术能力证明
            - 基于上述内容设计学术兴趣段落的亮点，并提出具体增强策略
            6.增强策略中的事实性与真实性原则：
            - 禁止在任何增强策略中过度编撰数据或杜撰不存在的经历细节
            - 所有增强策略必须基于素材表中提供的真实信息，不得虚构或夸大
            - 所有建议必须符合申请人现有的学术能力和专业水平
            - 使用"深化"时，应限于对现有经历的更详细描述，而非添加不存在的成就
            - 建议应具体、可行，申请人能够在现有基础上实际完成
            - 在没有足够信息支持的情况下，提供方向性建议而非具体事例
            7.科研和实习经历的严格事实性规则：
            - 禁止杜撰不存在的研究方法、数据、结果、成果或项目贡献
            - 禁止虚构与导师、同事或专家的交流内容
            - 禁止编造未在素材表中明确提及的专业术语、技术名称、研究工具或行业标准
            - 严禁为申请人添加未曾使用的专业软件、分析方法或研究技术
            - 所有研究和项目细节必须直接来源于素材表中的明确描述
            - 不得擅自提升项目规模、影响力或申请人的贡献程度
            - 对于素材表中模糊或笼统的描述，仅建议如何清晰表达，而非添加具体细节
            - 实习职责描述必须完全基于素材表提供的信息，不得扩展职责范围
            - 量化成果必须有素材表中的明确数据支持，禁止自行设定数值或比例

            分析流程
            第一步：专业方向分析
            1.判断用户输入类型：
            - 如果是明确的申请专业方向名称，必须以用户明确提供的申请方向为唯一分析对象，完全忽略素材表中的专业方向信息，不得对用户提供的方向名称进行任何修改或补充
            - 如果用户提供选校规划（包含多个院校及专业名称），严格执行以下分析：
            a. 分析表格中列出的所有院校专业名称
            b. 将专业按学科相似性归类为最多2个核心专业大类方向（如国际关系、公共政策等）
            c. 确保归类结果直接基于专业名称的学术内容，而非个人偏好
            d. 无论提交材料中包含什么其他信息，都必须以选校规划中的专业所属大类为唯一依据进行分类
            2.对每个识别出的专业方向进行招生倾向和评估重点分析

            第二步：段落素材策略与增强指南设计
            1.明确每个必要段落（专业兴趣塑造、学术基础展示、科研经历深化、实习经历深化、未来规划提升）的具体素材选择
            - 在专业兴趣段落中，必须只确定一个最合适的写作角度，不提供多个角度选择
            - 在科研经历和实习经历的"现有基础"描述后，必须插入一段"选择理由"，清晰阐述为何选择此经历素材而非其他经历，格式应为："相比[其他可选经历]，[所选经历]提供了[具体优势]，更能展示[能力/特质]"
            - 【严格规定】必须避免在不同段落之间重复使用相同的素材内容，每段落使用的核心素材必须是唯一的且不可重复；除非用户明确要求使用两个经历素材，此时可以在指定的段落中使用第二个经历素材，但仍需确保其在其他段落中不被重复使用
            - 【强制执行】在确定各段落素材前，必须先进行整体规划，确保每个重要经历只被分配到一个最合适的段落中使用；除非用户明确指出需要在特定段落中使用两个经历素材
            - 【明确禁止】同一个实习、研究或活动经历不得在多个段落中反复提及或使用
            - 【特别注意】专业兴趣段落与其他段落（尤其是科研经历和实习经历）必须使用完全不同的素材，如某经历已在专业兴趣段落中作为引发兴趣的因素提及，则不得再用于其他段落的主要素材
            - 在分配素材时，应优先为科研和实习段落保留最有力的相关经历，其他经历再用于其他段落
            2.对学术基础展示段落：
            - 如有成绩单transcript_analysis，必须详细分析其中的课程内容和表现
            - 重点应先确定与申请专业相符合的学术理论框架、学术能力及可迁移能力
            - 以这些能力和框架为主线，辅以具体的课程内容或课堂作业作为支撑例子
            - 避免以课程为主要亮点进行描述，而应以学术能力和专业相关技能为核心
            - 如成绩单缺失，则必须：
            a. 分析学生当前在读专业与申请专业的关联性和衔接点
            b. 挖掘素材表中的相关教育经历，包括课程项目、学术活动等
            c. 分析研究经历中的课堂作业部分，找出体现学术能力的亮点
            d. 基于以上信息构建学术基础展示段落
            - 重点挖掘与申请专业相关的关键课程、项目和学术能力
            - 避免直接讨论课程分数及GPA数值，而是分析展示学生在相关领域的学术优势和专业能力
            - 如成绩单中有特别突出或能体现专业相关能力的课程表现，必须重点分析其价值
            3.精确理解并执行用户的个性化定制要求：
            - 分析用户需求中的每一项指令（如"替换"、"优化"、"强化"等）
            - 对于"替换"类需求：提供全新内容，完全不使用原有素材
            - 对于"优化"类需求：提供详细具体的改写方案，包含明确示例 
            - 对于"增加"类需求：
            a. 首先判断需求类型：是要添加新的经历段落，还是在现有段落中添加多个经历
            b. 如需添加新段落：保持原段落完整，按相同格式创建新的经历段落（如"实习经历深化2"）
            c. 如需在现有段落内添加多个经历：在同一段落内使用子标题区分不同经历（如"经历1"、"经历2"）  
            d. 在两种情况下，都优先使用素材表中未使用的相关经历；仅在没有合适素材时才创建符合专业方向的新经历
            - 对于其他定制需求：根据具体指令提供相应的解决方案
            4.对于科研经历和实习经历，必须从学生提供的多个经历中选择一个最相关的具体经历
            5.在分配素材时，必须避免在专业兴趣段落与科研经历、实习经历段落之间重复使用核心素材；重点经历应优先保留给科研和实习段落
            6.审核完成的方案必须再次检查素材使用情况，确保没有素材在不同段落中重复出现
            7.如果发现素材重复使用，必须立即重新分配，保证每个段落使用的主要经历素材的唯一性
            8.为每个段落提供
            - 详细具体的现有基础分析
            - 针对用户个性化需求的明确解决方案
            - 可直接实施的详细增强策略
            9.增强策略的内容限制与科研实习经历特殊规定：      
            - 严格审核所有增强策略，确保不超出素材表提供的信息范围
            - 避免建议添加不切实际或需要过多专业知识才能完成的内容
            - 对于科研成果和实习成果，不得建议添加未实际完成的项目结果
            - 所有建议的量化成果必须基于素材表中已有的信息或合理推断
            - 增强策略应专注于深化现有素材的表达方式，而非大幅扩展内容范围
            - 科研经历深化部分：
            - 必须100%基于素材表中的具体描述
            - 禁止添加未在素材表中明确提及的研究方法、技术或发现
            - 避免过度专业化，不得自行添加专业术语或理论框架
            - 对于研究过程的建议仅限于清晰表达已有内容
            - 实习经历深化部分：
            - 必须严格遵循素材表中描述的职责范围和内容
            - 禁止扩大申请人的职责或添加未曾执行的任务
            - 量化成果必须有明确的事实依据，不得自行估算或创造数据
            - 不得虚构与同事、上级的互动或项目影响

            """,
            
            'consultant_role2': """
            # 角色
            我是一位资深的学术内容创作专家，曾为数百名学生成功撰写录取率极高的个人陈述。
            我拥有丰富的跨文化写作经验，精通中英双语表达，能够在保持文化真实性的同时创作符合英语思维和表达习惯的内容。
            我曾在多所国际顶尖大学的招生部门工作，深谙各学科领域的录取偏好和评判标准。
            我擅长将学生的素材转化为引人入胜的叙事，通过巧妙的结构安排和语言选择，使每位申请者的独特价值在众多申请中脱颖而出。

            """,
            
            'output_format2': """
            输出格式
            个人陈述（专业大类1：[专业名称]）
            专业兴趣塑造 [按照分析报告中的段落素材策略与增强指南组织内容，注重逻辑性，并且深入展开细节描述和观点叙述，减少素材的堆砌，注重描述的深度...] 
            学术基础展示 [按照分析报告中的段落素材策略与增强指南组织内容，突出相关课程的学习成果和技能提升，体现与该专业方向的契合...] 
            研究经历深化 [按照分析报告中的段落素材策略与增强指南组织内容，遵循STAR原则和总分总结构详细描述相关经历，与专业方向相联系...] 
            实习经历深化 [按照分析报告中的段落素材策略与增强指南组织内容，遵循STAR原则和总分总结构详细描述相关经历，与专业方向相联系...] 
            未来规划提升 [按照分析报告中的段落素材策略与增强指南组织内容，结合该专业方向提供具体且合理的规划...] 
            为何选择该专业和院校 [从专业发展前景、学术理念契合和个人成长等角度，逻辑性地阐述为何选择该专业领域深造，不针对具体学校] 
            结语 [简洁有力地总结申请者的优势、志向和对该专业的热情...]

            个人陈述（专业大类2：[专业名称]）
            [按相同结构组织第二个专业方向的个人陈述...]

            写作说明
            ●确保文章结构清晰，段落之间有良好的逻辑过渡
            ●所有非素材表中需要补充的内容必须保留中文并用【补充：】标记
            ●内容均使用纯中文表达，表述则符合英语思维习惯
            ●确保英文部分与中文补充内容之间的自然过渡
            ●技术术语和专业概念则使用准确的英文表达
            ●保持文章的整体连贯性和专业性
            ●重点突出申请者的优势，并与申请方向建立明确联系
            ●内容应真实可信，避免虚构经历或夸大成就
            ●每个主题部分应当是一个连贯的整体段落，而非多个松散段落
            ●"为何选择该专业和院校"部分应从专业角度进行逻辑论述，不针对具体学校
            ●在分析和撰写过程中充分参考成绩单中的具体课程表现，但不要体现任何分数
            ●确保内容精练，避免不必要的重复和冗余表达
            ●结语应简明扼要地总结全文，展现申请者的决心和愿景
            ●所有段落必须严格执行报告中的指导，不合并不同经历或添加额外内容
            ●避免出现"突然感兴趣"或"因此更感兴趣"等生硬转折，确保兴趣发展有合理的渐进过程
            ●各段落间应有内在的逻辑联系，而非简单罗列，每段内容应自然引出下一段内容
            ●确保经历与专业兴趣间的关联性具有说服力，展示清晰的思维发展路径
            ●在严格遵循分析报告的基础上，优化表述逻辑，确保内容流畅自然，避免生硬的连接和突兀的转折
            ●即使是按照报告中的指导进行写作，也应确保整体叙事具有内在一致性和合理的心理动机发展



            """,
            
            'consultant_task2': """
            任务描述
            1.基于提供的专业方向分析报告、成绩单和素材表，为指定的专业方向创作完整的个人陈述初稿
            2.严格按照分析报告中的"段落素材策略与增强指南"组织内容，确保每个段落都完全遵循报告中提供的所有个性化需求及增强策略点
            3.遵循STAR原则(情境-任务-行动-结果)呈现研究经历和实习经历，且只选择报告中指定的一个最相关经历
            4.突出申请者与申请方向的契合点，参考"招生倾向分析"部分的具体要点
            5.在正文中直接使用【补充：】标记所有非素材表中的内容
            6.确保段落间有自然过渡，保持文章整体连贯性
            7.所有段落内容必须严格遵循报告中对应的策略要点和增强策略，不添加未在报告中指定的内容
            8.科研经历和实习经历段落必须各自只选择一个最相关的经历，完全按照报告中指定的增强策略展开
            9.在遵循分析报告内容的同时，需要优化表述逻辑，确保内容之间的连贯性和自然过渡

            结构要求
            1.每个段落应采用"总-分-总"结构，第一句话承上启下，最后一句话总结该经历与目标专业的联系
            2.每个段落应有连贯性，作为一个整体呈现，避免过多分段，理想情况下每个主题部分应为一个连贯的段落
            3.控制整体字数，每个段落控制在150-200字左右，确保文书紧凑精炼
            4.第一段专业兴趣段落尤其需要注重逻辑性表述，避免堆砌过多素材
            5.增强句子之间的逻辑连接：
            - 确保每个新句子包含前一句子的关键词或概念
            - 使用指代词明确引用前文内容
            - 恰当使用过渡词和连接词
            - 建立清晰的因果关系，使用"因此"、"由此"、"正是"等词语明确前后句关系
            - 采用递进结构展示思想发展，从初始观察到深入思考，再到形成核心观点
            - 添加过渡句确保各点之间自然衔接，如"这种认识引导我..."、"通过这一探索..."- 确保每个段落形成完整的思想发展脉络，展现认知的深化过程   
            - 避免单纯并列不相关信息，而是通过逻辑词建立内在联系
            6.所有段落必须充分整合分析报告中的"个性化需求"和"增强策略"，同时保持简洁
            7.严格遵守报告中对每个段落的具体指导，不添加额外经历或合并不同经历
            8.强化内在逻辑推理，确保段落内容之间存在清晰的因果关系和思维发展脉络



            """
        }
        
        # 初始化 session_state 中的模板
        if 'templates' not in st.session_state:
            st.session_state.templates = self.default_templates.copy()

    def get_template(self, template_name: str) -> str:
        return st.session_state.templates.get(template_name, "")

    def update_template(self, template_name: str, new_content: str) -> None:
        st.session_state.templates[template_name] = new_content

    def reset_to_default(self):
        st.session_state.templates = self.default_templates.copy()

class TranscriptAnalyzer:
    def __init__(self, api_key: str, prompt_templates: PromptTemplates):
        self.prompt_templates = prompt_templates
        if 'templates' not in st.session_state:
            st.session_state.templates = self.prompt_templates.default_templates.copy()
            
        self.llm = ChatOpenAI(
            temperature=0.2,
            model=st.session_state.transcript_model,  # 使用session state中的模型
            api_key=api_key,
            base_url="https://openrouter.ai/api/v1",
            streaming=True
        )
    
    def extract_images_from_pdf(self, pdf_bytes):
        """从PDF中提取图像"""
        try:
            images = []
            pdf_document = fitz.open(stream=pdf_bytes, filetype="pdf")
            
            for page_num in range(len(pdf_document)):
                page = pdf_document[page_num]
                # 将页面直接转换为图像
                pix = page.get_pixmap(matrix=fitz.Matrix(2, 2))
                img_bytes = pix.tobytes("png")
                # 将图像编码为base64字符串
                img_base64 = base64.b64encode(img_bytes).decode('utf-8')
                images.append(img_base64)
            
            return images
        except Exception as e:
            logger.error(f"提取PDF图像时出错: {str(e)}")
            return []
    
    def analyze_transcript(self, pdf_bytes) -> Dict[str, Any]:
        try:
            if not hasattr(self, 'prompt_templates'):
                logger.error("prompt_templates not initialized")
                raise ValueError("Prompt templates not initialized properly")
            
            images = self.extract_images_from_pdf(pdf_bytes)
            if not images:
                return {
                    "status": "error",
                    "message": "无法从PDF中提取图像"
                }
            
            # 修改消息格式
            messages = [
                SystemMessage(content=self.prompt_templates.get_template('transcript_role')),
                HumanMessage(content=[  # 注意这里改成了列表
                    {
                        "type": "text",
                        "text": f"\n\n请分析这份成绩单，提取成绩信息，并以表格形式输出成绩信息。"
                    },
                    {
                        "type": "image_url",
                        "image_url": {
                            "url": f"data:image/png;base64,{images[0]}"
                        }
                    }
                ])
            ]
            
            # 创建一个队列用于流式输出
            message_queue = Queue()
            
            # 创建自定义回调处理器
            class QueueCallbackHandler(BaseCallbackHandler):
                def __init__(self, queue):
                    self.queue = queue
                    super().__init__()
                
                def on_llm_new_token(self, token: str, **kwargs) -> None:
                    self.queue.put(token)
            
            # 创建一个生成器函数，用于流式输出
            def token_generator():
                while True:
                    try:
                        token = message_queue.get(block=False)
                        yield token
                    except Empty:
                        if not thread.is_alive() and message_queue.empty():
                            break
                    time.sleep(0.01)
            
            # 在单独的线程中运行分析
            def run_analysis():
                try:
                    # 调用LLM进行分析
                    chain = LLMChain(llm=self.llm, prompt=ChatPromptTemplate.from_messages(messages))
                    result = chain.run(
                        {},
                        callbacks=[QueueCallbackHandler(message_queue)]
                    )
                    
                    message_queue.put("\n\n成绩单分析完成！")
                    thread.result = result
                    return result
                    
                except Exception as e:
                    message_queue.put(f"\n\n错误: {str(e)}")
                    logger.error(f"成绩单分析错误: {str(e)}")
                    thread.exception = e
                    raise e
            
            # 启动线程
            thread = Thread(target=run_analysis)
            thread.start()
            
            # 使用 st.write_stream 显示流式输出
            output_container = st.empty()
            with output_container:
                full_response = st.write_stream(token_generator())
            
            # 等待线程完成
            thread.join()
            
            # 获取结果
            if hasattr(thread, "exception") and thread.exception:
                raise thread.exception
            
            logger.info("成绩单分析完成")
            
            return {
                "status": "success",
                "transcript_analysis": full_response
            }
                
        except Exception as e:
            logger.error(f"成绩单分析错误: {str(e)}")
            return {
                "status": "error",
                "message": str(e)
            }
                

class BrainstormingAgent:
    def __init__(self, api_key: str, prompt_templates: PromptTemplates):
        self.llm = ChatOpenAI(
            temperature=0.2,
            model=st.session_state.text_model,  # 使用session state中的模型
            api_key=api_key,
            base_url="https://openrouter.ai/api/v1",
            streaming=True
        )
        self.prompt_templates = prompt_templates
        self.setup_chains()

    def setup_chains(self):
        # 背景分析 Chain
        strategist_prompt = ChatPromptTemplate.from_messages([
            ("system", f"{self.prompt_templates.get_template('consultant_role1')}\n\n"
                      f"任务:\n{self.prompt_templates.get_template('consultant_task1')}\n\n"
                      f"请按照以下格式输出:\n{self.prompt_templates.get_template('output_format1')}"),
            ("human", "选校规划school_plan：\n{school_plan}\n\n"
                     "成绩单transcript_analysis：\n{transcript_analysis}\n\n"
                     "个性化需求custom_requirements：\n{custom_requirements}\n\n"
                     "素材表document_content：\n{document_content}")
        ])
        
        self.strategist_chain = LLMChain(
            llm=self.llm,
            prompt=strategist_prompt,
            output_key="strategist_analysis",
            verbose=True
        )

        # 内容规划 Chain 
        creator_prompt = ChatPromptTemplate.from_messages([
            ("system", f"{self.prompt_templates.get_template('consultant_role2')}\n\n"
                      f"任务:\n{self.prompt_templates.get_template('consultant_task2')}\n\n"
                      f"请按照以下格式输出:\n{self.prompt_templates.get_template('output_format2')}"),
            ("human", "基于第一阶段的专业方向分析报告：\n{strategist_analysis}\n\n"
                     "成绩单transcript_analysis：\n{transcript_analysis}\n\n"
                     "素材表document_content：\n{document_content}\n\n"
                     "请创建详细的内容规划。")
        ])
        
        self.creator_chain = LLMChain(
            llm=self.llm,
            prompt=creator_prompt,
            output_key="creator_output",
            verbose=True
        )

    def process_strategist(self, document_content: str, school_plan: str, transcript_analysis: str = "", custom_requirements: str = "无定制需求") -> Dict[str, Any]:
        try:
            # 创建一个队列用于流式输出
            message_queue = Queue()
            
            # 创建自定义回调处理器，继承自 BaseCallbackHandler
            class QueueCallbackHandler(BaseCallbackHandler):
                def __init__(self, queue):
                    self.queue = queue
                    super().__init__()
                
                def on_llm_new_token(self, token: str, **kwargs) -> None:
                    self.queue.put(token)
            
            # 创建一个生成器函数，用于流式输出
            def token_generator():
                while True:
                    try:
                        token = message_queue.get(block=False)
                        yield token
                    except Empty:
                        if not thread.is_alive() and message_queue.empty():
                            break
                    time.sleep(0.01)
            
            # 在单独的线程中运行LLM
            def run_llm():
                try:
                    result = self.strategist_chain(
                        {
                            "document_content": document_content,
                            "school_plan": school_plan,
                            "transcript_analysis": transcript_analysis,
                            "custom_requirements": custom_requirements  # 添加默认值
                        },
                        callbacks=[QueueCallbackHandler(message_queue)]
                    )
                    # 将结果存储在线程对象中
                    thread.result = result
                    message_queue.put("\n\n分析完成！")
                    return result
                except Exception as e:
                    message_queue.put(f"\n\n错误: {str(e)}")
                    logger.error(f"Strategist processing error: {str(e)}")
                    thread.exception = e
                    raise e
            
            # 启动线程
            thread = Thread(target=run_llm)
            thread.start()
            
            # 使用 st.write_stream 显示流式输出
            output_container = st.empty()
            with output_container:
                full_response = st.write_stream(token_generator())
            
            # 等待线程完成
            thread.join()
            
            # 获取结果
            if hasattr(thread, "exception") and thread.exception:
                raise thread.exception
            
            logger.info("Strategist analysis completed successfully")
            
            # 从 full_response 中提取分析结果
            return {
                "status": "success",
                "strategist_analysis": full_response
            }
                
        except Exception as e:
            logger.error(f"Strategist processing error: {str(e)}")
            return {
                "status": "error",
                "message": str(e)
            }
    def process_creator(self, strategist_analysis: str, document_content: str, school_plan: str, transcript_analysis: str = "", custom_requirements: str = "无定制需求") -> Dict[str, Any]:
        try:
            # 创建一个队列用于流式输出
            message_queue = Queue()
            
            # 创建自定义回调处理器，继承自 BaseCallbackHandler
            class QueueCallbackHandler(BaseCallbackHandler):
                def __init__(self, queue):
                    self.queue = queue
                    super().__init__()
                
                def on_llm_new_token(self, token: str, **kwargs) -> None:
                    self.queue.put(token)
            
            # 创建一个生成器函数，用于流式输出
            def token_generator():
                while True:
                    try:
                        token = message_queue.get(block=False)
                        yield token
                    except Empty:
                        if not thread.is_alive() and message_queue.empty():
                            break
                    time.sleep(0.01)
            
            # 在单独的线程中运行LLM
            def run_llm():
                try:
                    result = self.creator_chain(
                        {
                            "strategist_analysis": strategist_analysis,
                            "document_content": document_content,  # 添加文档内容
                            "school_plan": school_plan,
                            "transcript_analysis": transcript_analysis,
                            "custom_requirements": custom_requirements
                        },
                        callbacks=[QueueCallbackHandler(message_queue)]
                    )
                    # 将结果存储在队列中
                    message_queue.put("\n\n规划完成！")
                    return result
                except Exception as e:
                    message_queue.put(f"\n\n错误: {str(e)}")
                    logger.error(f"Creator processing error: {str(e)}")
                    raise e
            
            # 启动线程
            thread = Thread(target=run_llm)
            thread.start()
            
            # 使用 st.write_stream 显示流式输出
            output_container = st.empty()
            with output_container:
                full_response = st.write_stream(token_generator())
            
            # 等待线程完成
            thread.join()
            
            # 获取结果
            if hasattr(thread, "_exception") and thread._exception:
                raise thread._exception
            
            logger.info("Creator analysis completed successfully")
            
            return {
                "status": "success",
                "creator_output": full_response
            }
                
        except Exception as e:
            logger.error(f"Creator processing error: {str(e)}")
            return {
                "status": "error",
                "message": str(e)
            }


def add_custom_css():
    st.markdown("""
    <style>
    /* 整体页面样式 */
    .main {
        padding: 2rem;
    }
    
    /* 标题样式 */
    h1, h2, h3 {
        color: #1e3a8a;
        font-weight: 600;
        margin-bottom: 1.5rem;
    }
    
    .page-title {
        text-align: center;
        font-size: 2.5rem;
        margin-bottom: 2rem;
        color: #1e3a8a;
        font-weight: bold;
        padding: 1rem;
        border-bottom: 3px solid #e5e7eb;
    }
    
    /* 文件上传区域样式 */
    .stFileUploader {
        margin-bottom: 2rem;
    }
    
    .stFileUploader > div > button {
        background-color: #f8fafc;
        color: #1e3a8a;
        border: 2px dashed #1e3a8a;
        border-radius: 8px;
        padding: 1rem;
        transition: all 0.3s ease;
    }
    
    .stFileUploader > div > button:hover {
        background-color: #f0f7ff;
        border-color: #2563eb;
    }
    
    /* 按钮样式 */
    .stButton > button {
        background-color: #1e3a8a;
        color: white;
        border-radius: 8px;
        padding: 0.75rem 1.5rem;
        font-weight: 500;
        border: none;
        width: 100%;
        transition: all 0.3s ease;
        box-shadow: 0 2px 4px rgba(0,0,0,0.1);
    }
    
    .stButton > button:hover {
        background-color: #2563eb;
        transform: translateY(-1px);
        box-shadow: 0 4px 6px rgba(0,0,0,0.1);
    }
    
    .stButton > button:disabled {
        background-color: #94a3b8;
        cursor: not-allowed;
    }
    
    /* 文本区域样式 */
    .stTextArea > div > div > textarea {
        border-radius: 8px;
        border: 1px solid #e2e8f0;
        padding: 0.75rem;
        font-size: 1rem;
        transition: all 0.3s ease;
    }
    
    .stTextArea > div > div > textarea:focus {
        border-color: #2563eb;
        box-shadow: 0 0 0 3px rgba(37,99,235,0.1);
    }
    
    /* 分析结果区域样式 */
    .analysis-container {
        background-color: white;
        border-radius: 12px;
        padding: 1.5rem;
        margin: 1rem 0;
        box-shadow: 0 4px 6px rgba(0,0,0,0.05);
        border: 1px solid #e5e7eb;
    }
    
    /* 成功消息样式 */
    .stSuccess {
        background-color: #ecfdf5;
        color: #065f46;
        padding: 1rem;
        border-radius: 8px;
        border-left: 4px solid #059669;
        margin: 1rem 0;
    }
    
    /* 错误消息样式 */
    .stError {
        background-color: #fef2f2;
        color: #991b1b;
        padding: 1rem;
        border-radius: 8px;
        border-left: 4px solid #dc2626;
        margin: 1rem 0;
    }
    
    /* 模型信息样式 */
    .model-info {
        background-color: #f0f7ff;
        padding: 0.75rem 1rem;
        border-radius: 8px;
        margin: 0.5rem 0;
        display: inline-block;
        font-size: 0.9rem;
        border: 1px solid #bfdbfe;
    }
    
    /* 双列布局样式 */
    .dual-column {
        display: flex;
        gap: 2rem;
        margin: 1rem 0;
    }
    
    .column {
        flex: 1;
        background-color: #f8fafc;
        padding: 1.5rem;
        border-radius: 12px;
        border: 1px solid #e5e7eb;
    }
    
    /* 分隔线样式 */
    hr {
        margin: 2rem 0;
        border: 0;
        height: 1px;
        background: linear-gradient(to right, transparent, #e5e7eb, transparent);
    }
    
    /* 标签页样式 */
    .stTabs {
        background-color: white;
        border-radius: 12px;
        padding: 1rem;
        box-shadow: 0 1px 3px rgba(0,0,0,0.1);
    }
    
    .stTab {
        padding: 1rem;
    }
    
    /* 展开器样式 */
    .streamlit-expanderHeader {
        background-color: #f8fafc;
        border-radius: 8px;
        padding: 0.75rem;
        font-weight: 500;
        color: #1e3a8a;
        border: 1px solid #e5e7eb;
    }
    
    .streamlit-expanderContent {
        background-color: white;
        border-radius: 0 0 8px 8px;
        padding: 1rem;
        border: 1px solid #e5e7eb;
        border-top: none;
    }
    
    /* 加载动画样式 */
    .stSpinner > div {
        border-color: #2563eb transparent transparent transparent;
    }
    
    /* 文档分析区域样式 */
    .doc-analysis-area {
        background-color: #ffffff;
        border-radius: 12px;
        padding: 1.5rem;
        margin: 1rem 0;
        box-shadow: 0 2px 4px rgba(0,0,0,0.05);
        border: 1px solid #e5e7eb;
    }
    
    .doc-analysis-area h3 {
        color: #1e3a8a;
        font-size: 1.25rem;
        margin-bottom: 1rem;
        padding-bottom: 0.5rem;
        border-bottom: 2px solid #e5e7eb;
    }
    
    /* 调整列宽度 */
    .column-adjust {
        padding: 0 1rem;
    }
    </style>
    """, unsafe_allow_html=True)


def read_docx(file_bytes):
    """读取 Word 文档内容，包括表格，并去除重复内容"""
    try:
        doc = Document(io.BytesIO(file_bytes))
        content_set = set()  # 用于存储已处理的内容，避免重复
        full_text = []
        
        # 读取普通段落
        for paragraph in doc.paragraphs:
            text = paragraph.text.strip()
            if text and text not in content_set:  # 只添加非空且未重复的内容
                content_set.add(text)
                full_text.append(text)
        
        # 读取表格内容
        for table in doc.tables:
            table_content = []
            header_row = []
            
            # 获取表头（第一行）
            if table.rows:
                for cell in table.rows[0].cells:
                    header_text = cell.text.strip()
                    if header_text:
                        header_row.append(header_text)
            
            # 处理表格内容（从第二行开始）
            for row_idx, row in enumerate(table.rows[1:], 1):
                row_content = {}
                for col_idx, cell in enumerate(row.cells):
                    if col_idx < len(header_row):  # 确保有对应的表头
                        cell_text = cell.text.strip()
                        if cell_text:
                            row_content[header_row[col_idx]] = cell_text
                
                if row_content:  # 只添加非空行
                    formatted_row = " | ".join([f"{header}: {value}" 
                                              for header, value in row_content.items()])
                    if formatted_row not in content_set:  # 避免重复内容
                        content_set.add(formatted_row)
                        table_content.append(formatted_row)
            
            if table_content:
                full_text.extend(table_content)
        
        # 使用换行符连接所有文本
        result = "\n".join(full_text)
        logger.info(f"成功读取文档内容，包含 {len(doc.tables)} 个表格")
        return result
    except Exception as e:
        logger.error(f"读取 Word 文档时出错: {str(e)}")
        return None

def initialize_session_state():
    if 'templates' not in st.session_state:
        prompt_templates = PromptTemplates()
        st.session_state.templates = prompt_templates.default_templates.copy()
    
    if 'prompt_templates' not in st.session_state:
        st.session_state.prompt_templates = PromptTemplates()

def main():
    initialize_session_state()
    
    langsmith_api_key = st.secrets["LANGCHAIN_API_KEY"]
    os.environ["LANGCHAIN_TRACING_V2"] = "true"
    os.environ["LANGCHAIN_API_KEY"] = langsmith_api_key
    os.environ["LANGCHAIN_PROJECT"] = "初稿脑暴平台"
    st.set_page_config(page_title="初稿脑暴助理平台", layout="wide")
    add_custom_css()
    st.markdown("<h1 class='page-title'>初稿脑暴助理</h1>", unsafe_allow_html=True)
    
    # 添加模型选择
    col1, col2 = st.columns(2)
    with col1:
        if 'transcript_model' not in st.session_state:
            st.session_state.transcript_model = st.secrets["TRANSCRIPT_MODEL"]
        
        transcript_model = st.selectbox(
            "选择图像分析模型",
            ["google/gemma-3-27b-it:free","google/gemini-2.0-flash-001","google/gemini-2.5-pro-exp-03-25:free", "qwen/qwen2.5-vl-32b-instruct:free","anthropic/claude-3.7-sonnet"],
            index=["google/gemma-3-27b-it:free","google/gemini-2.0-flash-001","google/gemini-2.5-pro-exp-03-25:free", "qwen/qwen2.5-vl-32b-instruct:free","anthropic/claude-3.7-sonnet"].index(st.session_state.transcript_model)
        )
        if transcript_model != st.session_state.transcript_model:
            st.session_state.transcript_model = transcript_model
            st.rerun()
            
    with col2:
        if 'text_model' not in st.session_state:
            st.session_state.text_model = st.secrets["OPENROUTER_MODEL"]
            
        text_model = st.selectbox(
            "选择文本分析模型",
            ["qwen/qwq-32b:free","qwen/qwq-32b","google/gemini-2.5-pro-exp-03-25:free", "deepseek/deepseek-chat-v3-0324:free", "deepseek/deepseek-r1:free","deepseek/deepseek-r1","anthropic/claude-3.7-sonnet"],
            index=["qwen/qwq-32b:free","qwen/qwq-32b","google/gemini-2.5-pro-exp-03-25:free", "deepseek/deepseek-chat-v3-0324:free", "deepseek/deepseek-r1:free","deepseek/deepseek-r1","anthropic/claude-3.7-sonnet"].index(st.session_state.text_model)
        )
        if text_model != st.session_state.text_model:
            st.session_state.text_model = text_model
            st.rerun()

    # 添加单个模型信息展开框
    with st.expander("查看模型详细信息", expanded=False):
        st.markdown("""
        ### 图像分析模型
        
        **google/gemma-3-27b-it:free**
        - 参数量：27B
        - Created Mar 12, 2025 | 96,000 context | $0/M  input tokens | $0\/M output tokens
        
        **google/gemini-2.0-flash-001**
        - Created Feb 5, 2025 |  1,000,000 context |  $0.1/M  input tokens  | $0.4\/M output tokens |  $0.0258\/K input imgs
        
        **google/gemini-2.5-pro-exp-03-25:free**
        - 最新版本的Gemini模型
        - Created Mar 25, 2025 | 1,000,000 context | $0/M  input tokens | $0\/M output tokens
        
        **qwen/qwen2.5-vl-32b-instruct:free**
        - 参数量：32B
        - Created Mar 24, 2025 | 8,192 context | $0/M  input tokens | $0\/M output tokens
                    
        **anthropic/claude-3.7-sonnet**
        - 参数量：32B
        - Created Feb 24, 2025 | 200,000 context | $3/M  input tokens | $15\/M output tokens | $4.8\/K input imgs
        
        ### 文本分析模型
        
        **qwen/qwq-32b:free**
        - 参数量：32B
        - Created Mar 5, 2025 | 40,000 context | $0/M input tokens | $0\/M output tokens
        
        **qwen/qwq-32b**
        - 完整版本，性能更优
        - Created Mar 5, 2025 | 131,072 context | $0.12/M input tokens  | $0.18\/M output tokens
        
        **google/gemini-2.5-pro-exp-03-25:free**
        - 最新版本的Gemini模型
        - Created Mar 25, 2025 | 1,000,000 context | $0/M input tokens | $0\/M output tokens
        
        **deepseek/deepseek-chat-v3-0324:free**
        - 最新的v3版本
        - Created Mar 24, 2025 | 131,072 context | $0/M input tokens | $0\/M output tokens
        
        **deepseek/deepseek-r1:free & deepseek/deepseek-r1**
        - Created Jan 20, 2025 | 163,840 context | $0/M input tokens | $0\/M output tokens
                    
        **anthropic/claude-3.7-sonnet**
        - 参数量：32B
        - Created Feb 24, 2025 | 200,000 context | $3/M input tokens | $15\/M output tokens | $4.8\/K input imgs
        
        """)

    # 确保在任何操作之前初始化 PromptTemplates
    if 'templates' not in st.session_state:
        prompt_templates = PromptTemplates()
        st.session_state.templates = prompt_templates.default_templates.copy()
    
    if 'prompt_templates' not in st.session_state:
        st.session_state.prompt_templates = PromptTemplates()
    
    tab1, tab2 = st.tabs(["初稿脑暴助理", "提示词设置"])
    
    # 初始化会话状态变量
    if 'document_content' not in st.session_state:
        st.session_state.document_content = None
    if 'transcript_file' not in st.session_state:
        st.session_state.transcript_file = None
    if 'transcript_analysis_done' not in st.session_state:
        st.session_state.transcript_analysis_done = False
    if 'transcript_analysis_result' not in st.session_state:
        st.session_state.transcript_analysis_result = None
    if 'strategist_analysis_done' not in st.session_state:
        st.session_state.strategist_analysis_done = False
    if 'strategist_analysis_result' not in st.session_state:
        st.session_state.strategist_analysis_result = None
    if 'creator_analysis_done' not in st.session_state:
        st.session_state.creator_analysis_done = False
    if 'creator_analysis_result' not in st.session_state:
        st.session_state.creator_analysis_result = None
    if 'show_transcript_analysis' not in st.session_state:
        st.session_state.show_transcript_analysis = False
    if 'show_strategist_analysis' not in st.session_state:
        st.session_state.show_strategist_analysis = False
    if 'show_creator_analysis' not in st.session_state:
        st.session_state.show_creator_analysis = False
    
    with tab1:
       
        
        transcript_file = st.file_uploader("上传成绩单（可选）", type=['pdf'])
        # 自动检查文件状态并清除相关内存
        if not transcript_file:
            st.session_state.transcript_file = None
            st.session_state.transcript_analysis_done = False
            st.session_state.transcript_analysis_result = None
            st.session_state.show_transcript_analysis = False
        
        # 添加分析成绩单按钮
        if st.button("分析成绩单", key="analyze_transcript", use_container_width=True):
            if transcript_file is not None:
                st.session_state.transcript_file = transcript_file
                st.session_state.show_transcript_analysis = True
                st.session_state.transcript_analysis_done = False
                st.rerun()
        
        # 修改文件上传部分，移除多文件支持
        
        
        uploaded_file = st.file_uploader("上传初稿文档", type=['docx'])  # 改为单文件上传
        
        # 处理上传的文件
        if uploaded_file:
            document_content = read_docx(uploaded_file.read())
            if document_content:
                st.session_state.document_content = document_content
                st.success(f"文件上传成功！")
                with st.expander("查看文档内容", expanded=False):
                    st.write(document_content)
            else:
                st.error("无法读取文件，请检查格式是否正确。")
        
        # 添加选校方案输入框
        school_plan = st.text_area(
            "选校方案",
            value="暂未选校",
            height=100,
            help="请输入已确定的选校方案，包括学校和专业信息"
        )
        
        # 添加自定义需求输入框
        custom_requirements = st.text_area(
            "定制需求（可选）",
            value="无定制需求",
            height=100,
            help="请输入特殊的定制需求，如果没有可以保持默认值"
        )
        
        # 修改按钮区域，只保留单文件模式
        button_col1, button_col2 = st.columns(2)
        with button_col1:
            if st.button("开始背景分析", key="start_analysis", use_container_width=True):
                if st.session_state.document_content:
                    st.session_state.show_strategist_analysis = True
                    st.session_state.strategist_analysis_done = False
                    st.session_state.creator_analysis_done = False
                    st.session_state.show_creator_analysis = False
                    st.rerun()
        
        with button_col2:
            continue_button = st.button(
                "继续内容规划", 
                key="continue_to_creator", 
                use_container_width=True
            )
            
            if continue_button:
                st.session_state.show_creator_analysis = True
                st.session_state.creator_analysis_done = False
                st.rerun()
                # 添加清除分析结果按钮
        if st.button("清除所有分析", key="clear_analysis", use_container_width=True):
            # 清除所有分析相关的session状态
            st.session_state.document_content = None
            st.session_state.strategist_analysis_done = False
            st.session_state.creator_analysis_done = False
            st.session_state.show_strategist_analysis = False
            st.session_state.show_creator_analysis = False
            st.success("✅ 所有分析结果已清除！")
            st.rerun()
        # 修改结果显示区域，只保留单文档逻辑
        results_container = st.container()
        
        # 显示成绩单分析（保持不变）
        if st.session_state.show_transcript_analysis:
            with results_container:
                st.markdown("---")
                st.subheader("📊 成绩单分析")
                
                if not st.session_state.transcript_analysis_done:
                    try:
                        # 确保 prompt_templates 存在
                        if 'prompt_templates' not in st.session_state:
                            st.session_state.prompt_templates = PromptTemplates()
                        
                        transcript_analyzer = TranscriptAnalyzer(
                            api_key=st.secrets["OPENROUTER_API_KEY"],  # 使用OpenRouter API密钥
                            prompt_templates=st.session_state.prompt_templates
                        )
                        
                        with st.spinner("正在分析成绩单..."):
                            # 处理成绩单分析
                            result = transcript_analyzer.analyze_transcript(
                                st.session_state.transcript_file
                            )
                            
                            if result["status"] == "success":
                                # 保存成绩单分析结果到 session_state
                                st.session_state.transcript_analysis_result = result["transcript_analysis"]
                                st.session_state.transcript_analysis_done = True
                                st.success("✅ 成绩单分析完成！")
                            else:
                                st.error(f"成绩单分析出错: {result['message']}")
                    
                    except Exception as e:
                        st.error(f"处理过程中出错: {str(e)}")
                else:
                    # 如果已经完成，直接显示结果
                    st.markdown(st.session_state.transcript_analysis_result)
                    st.success("✅ 成绩单分析完成！")
        
        # 修改背景分析显示，只保留单文档逻辑
        if st.session_state.show_strategist_analysis:
            with results_container:
                st.markdown("---")
                st.subheader("📊 第一阶段：背景分析")
                
                if not st.session_state.strategist_analysis_done:
                    try:
                        agent = BrainstormingAgent(
                            api_key=st.secrets["OPENROUTER_API_KEY"],
                            prompt_templates=st.session_state.prompt_templates
                        )
                        
                        with st.spinner("正在分析文档..."):
                            transcript_analysis = ""
                            if st.session_state.transcript_analysis_done:
                                transcript_analysis = st.session_state.transcript_analysis_result
                            
                            result = agent.process_strategist(
                                st.session_state.document_content,
                                school_plan,
                                transcript_analysis,
                                custom_requirements
                            )
                            
                            if result["status"] == "success":
                                st.session_state.strategist_analysis_result = result["strategist_analysis"]
                                st.session_state.strategist_analysis_done = True
                                st.success("✅ 背景分析完成！")
                            else:
                                st.error(f"背景分析出错: {result['message']}")
                    except Exception as e:
                        st.error(f"处理过程中出错: {str(e)}")
                else:
                    st.markdown(st.session_state.strategist_analysis_result)
                    st.success("✅ 背景分析完成！")
        
        # 修改内容规划显示，只保留单文档逻辑
        if st.session_state.show_creator_analysis:
            with results_container:
                st.markdown("---")
                st.subheader("📝 第二阶段：内容规划")
                
                if not st.session_state.creator_analysis_done:
                    try:
                        agent = BrainstormingAgent(
                            api_key=st.secrets["OPENROUTER_API_KEY"],
                            prompt_templates=st.session_state.prompt_templates
                        )
                        
                        with st.spinner("正在规划内容..."):
                            creator_result = agent.process_creator(
                                strategist_analysis=st.session_state.strategist_analysis_result,
                                document_content=st.session_state.document_content,
                                school_plan=school_plan,
                                transcript_analysis=st.session_state.transcript_analysis_result,
                                custom_requirements=custom_requirements
                            )
                            
                            if creator_result["status"] == "success":
                                st.session_state.creator_analysis_result = creator_result["creator_output"]
                                st.session_state.creator_analysis_done = True
                                st.success("✅ 内容规划完成！")
                            else:
                                st.error(f"内容规划出错: {creator_result['message']}")
                    except Exception as e:
                        st.error(f"处理过程中出错: {str(e)}")
                else:
                    st.markdown(st.session_state.creator_analysis_result)
                    st.success("✅ 内容规划完成！")
            # 修改模型信息显示
        st.markdown(f"<div class='model-info'>🤖 图像分析当前使用模型: <b>{st.session_state.transcript_model}</b></div>", unsafe_allow_html=True)
        st.markdown(f"<div class='model-info'>🤖 背景分析及内容规划当前使用模型: <b>{st.session_state.text_model}</b></div>", unsafe_allow_html=True)
        
    
    with tab2:
        st.title("提示词设置")
        
        prompt_templates = st.session_state.prompt_templates
        
        # Agent 1 设置
        st.subheader("Agent 1 - 档案策略师")
        consultant_role1 = st.text_area(
            "角色设定",
            value=prompt_templates.get_template('consultant_role1'),
            height=200,
            key="consultant_role1"
        )
        
        consultant_task1 = st.text_area(
            "任务说明",
            value=prompt_templates.get_template('consultant_task1'),
            height=200,
            key="consultant_task1"
        )

        output_format1 = st.text_area(
            "输出格式",
            value=prompt_templates.get_template('output_format1'),
            height=200,
            key="output_format1"
        )
        # Agent 2 设置
        st.subheader("Agent 2 - 内容创作师")
        consultant_role2 = st.text_area(
            "角色设定",
            value=prompt_templates.get_template('consultant_role2'),
            height=200,
            key="consultant_role2"
        )

        consultant_task2 = st.text_area(
            "任务说明",
            value=prompt_templates.get_template('consultant_task2'),
            height=200,
            key="consultant_task2"
        )

        output_format2 = st.text_area(
            "输出格式",
            value=prompt_templates.get_template('output_format2'),
            height=200,
            key="output_format2"
        )
        
        col1, col2 = st.columns(2)
        with col1:
            if st.button("更新提示词", key="update_prompts"):
                prompt_templates.update_template('consultant_role1', consultant_role1)
                prompt_templates.update_template('output_format1', output_format1)
                prompt_templates.update_template('consultant_task1', consultant_task1)
                prompt_templates.update_template('consultant_role2', consultant_role2)
                prompt_templates.update_template('output_format2', output_format2)
                prompt_templates.update_template('consultant_task2', consultant_task2)
                st.success("✅ 提示词已更新！")
        
        with col2:
            if st.button("重置为默认提示词", key="reset_prompts"):
                prompt_templates.reset_to_default()
                st.rerun()

if __name__ == "__main__":
    main()